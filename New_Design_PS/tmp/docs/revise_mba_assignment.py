from pathlib import Path
import re
import tempfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


INPUT_PATH = Path(r"C:\Users\aiuser\Desktop\OMB7001_Organizational Analysis of Intragen.docx")
OUTPUT_PATH = Path(r"c:\Users\aiuser\Documents\GitHub\powershell-xml-prod\New_Design_PS\output\doc\OMB7001_Organizational Analysis of Intragen_revised.docx")


def insert_paragraph_after(paragraph, text, style_name):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_paragraph.style = style_name
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_paragraph_before(paragraph, text, style_name):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_paragraph.style = style_name
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_paragraph_after_table(table, text, style_name):
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_paragraph = Paragraph(new_p, table._parent)
    if style_name:
        new_paragraph.style = style_name
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_table_after_paragraph(doc, paragraph, rows, cols, style_name=None):
    table = doc.add_table(rows=rows, cols=cols)
    if style_name:
        table.style = style_name
    paragraph._p.addnext(table._tbl)
    return table


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def find_paragraph(doc, prefix, style_name=None):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if style_name and paragraph.style and paragraph.style.name != style_name:
            continue
        if text.startswith(prefix):
            return paragraph
    style_suffix = f" with style {style_name!r}" if style_name else ""
    raise ValueError(f"Paragraph starting with {prefix!r}{style_suffix} not found")


def insert_after_prefix(doc, prefix, texts):
    anchor = find_paragraph(doc, prefix)
    last = anchor
    style_name = anchor.style.name if anchor.style else None
    for text in texts:
        last = insert_paragraph_after(last, text, style_name)


def replace_bibliography(doc, bibliography_entries):
    bibliography_heading = find_paragraph(doc, "Βιβλιογραφία", style_name="Heading 1")
    appendices_heading = find_paragraph(doc, "Παραρτήματα", style_name="Heading 1")

    deleting = False
    to_delete = []
    for paragraph in doc.paragraphs:
        if paragraph._p is bibliography_heading._p:
            deleting = True
            continue
        if paragraph._p is appendices_heading._p:
            break
        if deleting:
            to_delete.append(paragraph)

    for paragraph in reversed(to_delete):
        delete_paragraph(paragraph)

    last = bibliography_heading
    style_name = "Normal"
    for entry in bibliography_entries:
        last = insert_paragraph_after(last, entry, style_name)


def replace_text_in_paragraphs(doc, replacements):
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        if new_text != text:
            paragraph.text = new_text


def set_paragraph_text_by_prefix(doc, prefix, new_text, style_name=None):
    paragraph = find_paragraph(doc, prefix, style_name=style_name)
    paragraph.text = new_text


def delete_paragraphs_between(doc, start_paragraph, end_paragraph):
    deleting = False
    to_delete = []
    for paragraph in doc.paragraphs:
        if paragraph._p is start_paragraph._p:
            deleting = True
            continue
        if paragraph._p is end_paragraph._p:
            break
        if deleting:
            to_delete.append(paragraph)

    for paragraph in reversed(to_delete):
        delete_paragraph(paragraph)


def replace_page_refs_in_body(doc):
    bibliography_heading = find_paragraph(doc, "Βιβλιογραφία", style_name="Heading 1")
    for paragraph in doc.paragraphs:
        if paragraph._p is bibliography_heading._p:
            break
        text = paragraph.text
        new_text = re.sub(r",\s*pp?\.\s*", ", σελ. ", text)
        if new_text != text:
            paragraph.text = new_text


def polish_cover_page(doc):
    cover_updates = {
        10: "Organizational Analysis of Intragen:",
        11: "Structure, Culture and Improvement Recommendations",
        14: "Programme Name: Master of Business Administration",
        15: "Module Name and Code: Management and Organisational Analysis (OMB7001)",
        16: "Student ID Number: 2517563",
        17: "Lecturer's Name: Tzakostas Panagiotis",
        18: "Word Count: 5144 (excluding table of contents, bibliography and appendices)",
    }
    for index, text in cover_updates.items():
        paragraph = doc.paragraphs[index]
        paragraph.text = text
        if index in (10, 11):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def sync_manual_contents(doc):
    contents_heading = find_paragraph_exact(doc, "Περιεχόμενα", style_name="Heading 1")
    first_body_heading = find_paragraph_exact(doc, "1. Εισαγωγή", style_name="Heading 1")

    toc_entries = []
    collecting = False
    for paragraph in doc.paragraphs:
        if paragraph._p is contents_heading._p:
            collecting = True
            continue
        if paragraph._p is first_body_heading._p:
            break
        if collecting:
            toc_entries.append(paragraph)

    new_entries = [
        "1. Εισαγωγή\t3",
        "2. Θεωρητικό πλαίσιο\t3",
        "2.1 Οργανωσιακή δομή\t3",
        "2.2 Οργανωσιακή κουλτούρα\t4",
        "2.3 Οργανωσιακή μάθηση και οργανισμός που μαθαίνει\t5",
        "3. Μεθοδολογία\t6",
        "4. Οργανωσιακή Ανάλυση\t7",
        "4.1 Παρουσίαση της εταιρείας\t7",
        "4.2 Συνολική οργανωσιακή δομή και κατανομή ανθρώπινου δυναμικού\t8",
        "4.3 Γενική ερμηνεία της δομής της Intragen\t8",
        "4.4 Βασικά χαρακτηριστικά του οργανωσιακού περιβάλλοντος\t9",
        "5. Ανάλυση βασικών οργανωσιακών θεμάτων\t9",
        "5.1 Μάθηση και Ανάπτυξη\t10",
        "5.2 Προσαρμοστικότητα και Καινοτομία\t12",
        "5.3 Κίνητρα και Δέσμευση Εργαζομένων\t15",
        "6. Προτάσεις Βελτίωσης\t17",
        "6.1 Δημιουργία ομάδας ανάπτυξης εσωτερικών εφαρμογών και εργαλείων παραγωγικότητας\t18",
        "6.2 Πιο δομημένη μετάβαση από την εκπαίδευση στην ενεργή συμμετοχή σε έργα\t19",
        "6.3 Ενίσχυση του αναπτυξιακού ρόλου του επικεφαλής ομάδας και καλύτερη αξιοποίηση της αξιολόγησης\t21",
        "7. Συμπέρασμα\t22",
        "Βιβλιογραφία\t23",
        "Παραρτήματα\t26",
        "Παράρτημα Α: Συμπληρωμένο OCAI\t26",
        "Παράρτημα Β: Συμπληρωμένο ερωτηματολόγιο οργανωσιακών τύπων του Mintzberg\t32",
    ]

    for paragraph, text in zip(toc_entries, new_entries):
        paragraph.text = text
        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.first_line_indent = Inches(0)
        if re.match(r"^\d+\.\d+\s", text):
            paragraph.paragraph_format.left_indent = Inches(0.35)
        else:
            paragraph.paragraph_format.left_indent = Inches(0)

    for paragraph in reversed(toc_entries[len(new_entries):]):
        delete_paragraph(paragraph)

    first_body_heading.paragraph_format.page_break_before = True


def normalize_text_paragraphs(doc):
    first_body_heading = find_paragraph_exact(doc, "1. Εισαγωγή", style_name="Heading 1")
    first_body_index = next(
        i for i, paragraph in enumerate(doc.paragraphs) if paragraph._p is first_body_heading._p
    )

    for index, paragraph in enumerate(doc.paragraphs):
        if index < first_body_index:
            continue
        if paragraph_has_drawing(paragraph):
            continue

        text = paragraph.text
        new_text = text
        new_text = re.sub(r"\[([^\]]*\d{4}[a-z]?(?:[^\]]*)?)\]", r"(\1)", new_text)
        new_text = new_text.replace("–", "-").replace("—", "-").replace("‑", "-")
        new_text = new_text.replace("’", "'").replace("‘", "'")
        new_text = re.sub(r" {2,}", " ", new_text)
        new_text = re.sub(r"\(\s+", "(", new_text)
        new_text = re.sub(r"\s+\)", ")", new_text)

        if new_text != text:
            paragraph.text = new_text


def apply_body_spacing(doc):
    bibliography_heading = find_paragraph(doc, "Βιβλιογραφία", style_name="Heading 1")
    for paragraph in doc.paragraphs:
        if paragraph._p is bibliography_heading._p:
            break
        if paragraph.style and paragraph.style.name == "Normal":
            paragraph.paragraph_format.line_spacing = 1.5

    appendix_b_heading = find_paragraph_exact(
        doc,
        "Παράρτημα Β: Συμπληρωμένο ερωτηματολόγιο οργανωσιακών τύπων του Mintzberg",
        style_name="Heading 1",
    )
    appendix_b_index = next(
        i for i, paragraph in enumerate(doc.paragraphs) if paragraph._p is appendix_b_heading._p
    )
    for paragraph in doc.paragraphs[appendix_b_index:]:
        if paragraph.style and paragraph.style.name == "Normal":
            paragraph.paragraph_format.line_spacing = 1.5


def paragraph_has_drawing(paragraph):
    return bool(paragraph._element.xpath('.//*[local-name()="drawing" or local-name()="pict"]'))


def find_paragraph_exact(doc, text, style_name=None):
    for paragraph in doc.paragraphs:
        if style_name and paragraph.style and paragraph.style.name != style_name:
            continue
        if paragraph.text.strip() == text:
            return paragraph
    style_suffix = f" with style {style_name!r}" if style_name else ""
    raise ValueError(f"Paragraph exactly matching {text!r}{style_suffix} not found")


def find_last_paragraph_exact(doc, text, style_name=None):
    for paragraph in reversed(doc.paragraphs):
        if style_name and paragraph.style and paragraph.style.name != style_name:
            continue
        if paragraph.text.strip() == text:
            return paragraph
    style_suffix = f" with style {style_name!r}" if style_name else ""
    raise ValueError(f"Last paragraph exactly matching {text!r}{style_suffix} not found")


def find_last_paragraph_exact_before(doc, before_text, target_text, style_name=None):
    before_paragraph = find_paragraph_exact(doc, before_text, style_name=style_name)
    found = None
    for paragraph in doc.paragraphs:
        if paragraph._p is before_paragraph._p:
            break
        if paragraph.text.strip() == target_text:
            found = paragraph
    if found is None:
        style_suffix = f" with style {style_name!r}" if style_name else ""
        raise ValueError(
            f"Paragraph matching {target_text!r} before {before_text!r}{style_suffix} not found"
        )
    return found


def find_next_drawing_paragraph(doc, anchor_paragraph):
    paragraphs = doc.paragraphs
    start = next(
        i for i, paragraph in enumerate(paragraphs) if paragraph._element is anchor_paragraph._element
    ) + 1
    for paragraph in paragraphs[start:]:
        if paragraph_has_drawing(paragraph):
            return paragraph
    raise ValueError(f"No drawing paragraph found after {anchor_paragraph.text!r}")


def extract_first_image_from_paragraph(paragraph):
    blips = paragraph._element.xpath('.//*[local-name()="blip"]')
    for blip in blips:
        rel_id = blip.get(qn("r:embed"))
        if not rel_id:
            continue
        part = paragraph.part.related_parts[rel_id]
        suffix = Path(str(part.partname)).suffix or ".png"
        extent_nodes = paragraph._element.xpath('.//*[local-name()="extent"]')
        width_inches = None
        if extent_nodes:
            cx = extent_nodes[0].get("cx")
            if cx and cx.isdigit():
                width_inches = int(cx) / 914400
        return part.blob, suffix, width_inches
    raise ValueError("No embedded image found in paragraph")


def load_chart_font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_centered_multiline_text(draw, center_x, top_y, text, font, fill, spacing=6):
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    draw.multiline_text(
        (center_x - text_width / 2, top_y),
        text,
        font=font,
        fill=fill,
        spacing=spacing,
        align="center",
    )
    return text_width, text_height


def build_ocai_dimension_chart(title, series_labels, category_labels, values_by_series):
    width, height = 1500, 800
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = load_chart_font(30, bold=True)
    axis_font = load_chart_font(20)
    tick_font = load_chart_font(16)
    value_font = load_chart_font(15, bold=True)
    legend_title_font = load_chart_font(18, bold=True)
    legend_font = load_chart_font(18)

    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((width - (title_bbox[2] - title_bbox[0])) / 2, 15), title, font=title_font, fill="black")

    left_margin = 100
    right_margin = 235
    top_margin = 80
    bottom_margin = 680
    plot_width = width - left_margin - right_margin
    plot_height = bottom_margin - top_margin
    y_max = 50
    y_step = 10

    grid_color = (220, 220, 220)
    axis_color = (0, 0, 0)
    bar_colors = [
        (31, 119, 180),
        (255, 127, 14),
        (44, 160, 44),
        (214, 39, 40),
    ]

    for tick in range(0, y_max + y_step, y_step):
        y = bottom_margin - (tick / y_max) * plot_height
        draw.line((left_margin, y, width - right_margin, y), fill=grid_color, width=1)
        tick_bbox = draw.textbbox((0, 0), str(tick), font=tick_font)
        draw.text((left_margin - 18 - (tick_bbox[2] - tick_bbox[0]), y - 10), str(tick), font=tick_font, fill="black")

    draw.line((left_margin, top_margin, left_margin, bottom_margin), fill=axis_color, width=2)
    draw.line((left_margin, bottom_margin, width - right_margin, bottom_margin), fill=axis_color, width=2)

    y_label = "Βαθμολογία"
    y_label_image = Image.new("RGBA", (220, 60), (255, 255, 255, 0))
    y_label_draw = ImageDraw.Draw(y_label_image)
    y_label_bbox = y_label_draw.textbbox((0, 0), y_label, font=axis_font)
    y_label_draw.text((0, 0), y_label, font=axis_font, fill="black")
    rotated_y_label = y_label_image.crop((0, 0, y_label_bbox[2] - y_label_bbox[0], y_label_bbox[3] - y_label_bbox[1])).rotate(90, expand=True)
    image.paste(rotated_y_label, (20, top_margin + (plot_height - rotated_y_label.size[1]) // 2), rotated_y_label)

    group_count = len(category_labels)
    series_count = len(series_labels)
    group_width = plot_width / group_count
    bar_width = min(48, group_width / (series_count + 1.6))

    for category_index, label in enumerate(category_labels):
        center_x = left_margin + (category_index + 0.5) * group_width
        start_x = center_x - (series_count * bar_width) / 2
        for series_index, series_values in enumerate(values_by_series):
            value = series_values[category_index]
            x1 = start_x + series_index * bar_width
            x2 = x1 + bar_width - 4
            y1 = bottom_margin - (value / y_max) * plot_height
            draw.rectangle((x1, y1, x2, bottom_margin), fill=bar_colors[series_index])

            value_text = f"{value:g}"
            value_bbox = draw.textbbox((0, 0), value_text, font=value_font)
            draw.text(
                (x1 + ((x2 - x1) - (value_bbox[2] - value_bbox[0])) / 2, y1 - 26),
                value_text,
                font=value_font,
                fill="black",
            )

        draw_centered_multiline_text(draw, center_x, bottom_margin + 18, label, axis_font, "black")

    legend_x = width - right_margin + 35
    legend_y = top_margin + 10
    draw.rounded_rectangle((legend_x, legend_y, width - 25, legend_y + 180), radius=10, outline=(180, 180, 180), width=2)
    legend_title = "Τύπος κουλτούρας"
    draw.text((legend_x + 14, legend_y + 12), legend_title, font=legend_title_font, fill="black")

    for index, label in enumerate(series_labels):
        row_y = legend_y + 48 + index * 30
        draw.rectangle((legend_x + 14, row_y + 4, legend_x + 54, row_y + 22), fill=bar_colors[index])
        draw.text((legend_x + 66, row_y), label, font=legend_font, fill="black")

    return image


def build_ocai_profile_chart(title, category_labels, values):
    width, height = 1500, 800
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = load_chart_font(30, bold=True)
    axis_font = load_chart_font(20)
    tick_font = load_chart_font(16)
    value_font = load_chart_font(16, bold=True)

    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((width - (title_bbox[2] - title_bbox[0])) / 2, 15), title, font=title_font, fill="black")

    left_margin = 110
    right_margin = 60
    top_margin = 85
    bottom_margin = 680
    plot_width = width - left_margin - right_margin
    plot_height = bottom_margin - top_margin
    y_max = 45
    y_step = 5

    grid_color = (220, 220, 220)
    bar_color = (31, 119, 180)

    for tick in range(0, y_max + y_step, y_step):
        y = bottom_margin - (tick / y_max) * plot_height
        draw.line((left_margin, y, width - right_margin, y), fill=grid_color, width=1)
        tick_bbox = draw.textbbox((0, 0), str(tick), font=tick_font)
        draw.text((left_margin - 18 - (tick_bbox[2] - tick_bbox[0]), y - 10), str(tick), font=tick_font, fill="black")

    draw.line((left_margin, top_margin, left_margin, bottom_margin), fill="black", width=2)
    draw.line((left_margin, bottom_margin, width - right_margin, bottom_margin), fill="black", width=2)

    y_label = "Μέσος όρος"
    y_label_image = Image.new("RGBA", (220, 60), (255, 255, 255, 0))
    y_label_draw = ImageDraw.Draw(y_label_image)
    y_label_bbox = y_label_draw.textbbox((0, 0), y_label, font=axis_font)
    y_label_draw.text((0, 0), y_label, font=axis_font, fill="black")
    rotated_y_label = y_label_image.crop((0, 0, y_label_bbox[2] - y_label_bbox[0], y_label_bbox[3] - y_label_bbox[1])).rotate(90, expand=True)
    image.paste(rotated_y_label, (20, top_margin + (plot_height - rotated_y_label.size[1]) // 2), rotated_y_label)

    category_count = len(category_labels)
    group_width = plot_width / category_count
    bar_width = min(270, group_width * 0.8)

    for index, (label, value) in enumerate(zip(category_labels, values)):
        center_x = left_margin + (index + 0.5) * group_width
        x1 = center_x - bar_width / 2
        x2 = center_x + bar_width / 2
        y1 = bottom_margin - (value / y_max) * plot_height
        draw.rectangle((x1, y1, x2, bottom_margin), fill=bar_color)

        value_text = f"{value:.1f}"
        value_bbox = draw.textbbox((0, 0), value_text, font=value_font)
        draw.text(
            (center_x - (value_bbox[2] - value_bbox[0]) / 2, y1 - 28),
            value_text,
            font=value_font,
            fill="black",
        )
        draw_centered_multiline_text(draw, center_x, bottom_margin + 18, label, axis_font, "black")

    return image


def replace_image_before_caption(doc, caption_text, image):
    caption = find_paragraph_exact(doc, caption_text)
    paragraphs = doc.paragraphs
    caption_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)

    width_inches = 6.8
    drawing_paragraph = None
    if caption_index > 0 and paragraph_has_drawing(paragraphs[caption_index - 1]):
        drawing_paragraph = paragraphs[caption_index - 1]

    if drawing_paragraph is not None:
        _, _, existing_width_inches = extract_first_image_from_paragraph(drawing_paragraph)
        if existing_width_inches:
            width_inches = existing_width_inches
        delete_paragraph(drawing_paragraph)

    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            temp_path = Path(tmp.name)
        image.save(str(temp_path), format="PNG")

        picture_paragraph = insert_paragraph_before(caption, "", "Normal")
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = picture_paragraph.add_run()
        run.add_picture(str(temp_path), width=Inches(width_inches))
    finally:
        if temp_path and temp_path.exists():
            temp_path.unlink()


def replace_ocai_profile_figures(doc):
    category_labels = ["Συνεργατική", "Καινοτομική", "Αγορακεντρική", "Ιεραρχική"]

    current_profile = build_ocai_profile_chart(
        "Προφίλ οργανωσιακής κουλτούρας - Υφιστάμενη κατάσταση",
        category_labels,
        [28.3, 29.2, 25.0, 17.5],
    )
    desired_profile = build_ocai_profile_chart(
        "Προφίλ οργανωσιακής κουλτούρας - Επιθυμητή κατάσταση",
        category_labels,
        [33.3, 34.2, 18.3, 14.2],
    )

    replace_image_before_caption(
        doc,
        "Σχήμα 3. Προφίλ της υφιστάμενης οργανωσιακής κουλτούρας της εταιρείας",
        current_profile,
    )
    replace_image_before_caption(
        doc,
        "Σχήμα 5. Προφίλ της επιθυμητής οργανωσιακής κουλτούρας της εταιρείας",
        desired_profile,
    )


def replace_ocai_dimension_figures(doc):
    series_labels = ["Συνεργατική", "Καινοτομική", "Αγορακεντρική", "Ιεραρχική"]
    category_labels = [
        "Κυρίαρχα\nχαρακτηριστικά",
        "Οργανωσιακή\nηγεσία",
        "Διοίκηση\nεργαζομένων",
        "Στοιχείο\nσυνοχής",
        "Στρατηγικές\nεμφάσεις",
        "Κριτήρια\nεπιτυχίας",
    ]

    current_values = [
        [30, 30, 35, 30, 20, 25],
        [30, 25, 25, 30, 35, 30],
        [25, 25, 20, 20, 30, 30],
        [15, 20, 20, 20, 15, 15],
    ]
    desired_values = [
        [35, 40, 40, 35, 35, 40],
        [35, 25, 20, 35, 35, 30],
        [10, 10, 10, 10, 10, 10],
        [20, 25, 30, 20, 20, 20],
    ]

    current_chart = build_ocai_dimension_chart(
        "Αποτελέσματα οργανωσιακής κουλτούρας ανά διάσταση - Υφιστάμενη κατάσταση",
        series_labels,
        category_labels,
        current_values,
    )
    desired_chart = build_ocai_dimension_chart(
        "Αποτελέσματα οργανωσιακής κουλτούρας ανά διάσταση - Επιθυμητή κατάσταση",
        series_labels,
        category_labels,
        desired_values,
    )

    replace_image_before_caption(
        doc,
        "Σχήμα 4. Αναλυτική αποτύπωση των αποτελεσμάτων για την υφιστάμενη οργανωσιακή κουλτούρα",
        current_chart,
    )
    replace_image_before_caption(
        doc,
        "Σχήμα 6. Αναλυτική αποτύπωση των αποτελεσμάτων για την επιθυμητή οργανωσιακή κουλτούρα",
        desired_chart,
    )


def restore_body_figure_1(doc):
    caption_text = "Σχήμα 1. Κατανομή του ανθρώπινου δυναμικού της εταιρείας ανά βασικό τμήμα"
    caption = find_paragraph_exact(doc, caption_text)
    paragraphs = doc.paragraphs
    caption_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    if caption_index > 0 and paragraph_has_drawing(paragraphs[caption_index - 1]):
        return

    original_doc = Document(str(INPUT_PATH))
    original_caption = find_paragraph_exact(original_doc, caption_text)
    original_paragraphs = original_doc.paragraphs
    original_caption_index = next(
        i for i, paragraph in enumerate(original_paragraphs) if paragraph._p is original_caption._p
    )
    source_paragraph = None
    for paragraph in reversed(original_paragraphs[:original_caption_index]):
        if paragraph_has_drawing(paragraph):
            source_paragraph = paragraph
            break
    if source_paragraph is None:
        raise ValueError("Original Figure 1 image paragraph not found")

    blob, suffix, width_inches = extract_first_image_from_paragraph(source_paragraph)
    picture_paragraph = insert_paragraph_before(caption, "", "Normal")
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(blob)
            temp_path = Path(tmp.name)
        run = picture_paragraph.add_run()
        if width_inches:
            run.add_picture(str(temp_path), width=Inches(width_inches))
        else:
            run.add_picture(str(temp_path), width=Inches(2.45))
    finally:
        if temp_path and temp_path.exists():
            temp_path.unlink()


def insert_ocai_comparison_table(doc):
    source_before_recommendations = find_last_paragraph_exact_before(
        doc,
        "6.1 Δημιουργία ομάδας ανάπτυξης εσωτερικών εφαρμογών και εργαλείων παραγωγικότητας",
        "Πηγή: Συμπληρωμένο ερωτηματολόγιο οργανωσιακής κουλτούρας και δημιουργία του συγγραφέα.",
    )

    comparison_table = insert_table_after_paragraph(doc, source_before_recommendations, 5, 4, "Normal Table")
    comparison_rows = [
        ["Τύπος κουλτούρας", "Τρέχουσα", "Επιθυμητή", "Μεταβολή"],
        ["Συνεργατική (Clan)", "28.3", "33.3", "+5.0"],
        ["Καινοτομική (Adhocracy)", "29.2", "34.2", "+5.0"],
        ["Αγορακεντρική (Market)", "25.0", "18.3", "-6.7"],
        ["Ιεραρχική (Hierarchy)", "17.5", "14.2", "-3.3"],
    ]
    for row_idx, row_values in enumerate(comparison_rows):
        for col_idx, value in enumerate(row_values):
            comparison_table.cell(row_idx, col_idx).text = value
    format_table(comparison_table)

    caption = insert_paragraph_after_table(
        comparison_table,
        "Πίνακας 1. Συνθετική σύγκριση υφιστάμενης και επιθυμητής οργανωσιακής κουλτούρας βάσει OCAI",
        "Appendix Caption",
    )
    source = insert_paragraph_after(
        caption,
        "",
        "Appendix Source",
    )
    apply_source_text(
        source,
        "Πηγή: Συμπληρωμένο ερωτηματολόγιο OCAI της παρούσας μελέτης και επεξεργασία του συγγραφέα.",
    )


def style_exists(doc, style_name):
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


def ensure_appendix_styles(doc):
    styles = doc.styles

    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        styles[style_name].font.name = "Times New Roman"

    if not style_exists(doc, "Appendix Page Heading"):
        style = styles.add_style("Appendix Page Heading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Heading 2"]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(12)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if not style_exists(doc, "Appendix Section Heading"):
        style = styles.add_style("Appendix Section Heading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Heading 3"]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(11)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True

    if not style_exists(doc, "Appendix Caption"):
        style = styles.add_style("Appendix Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(10.5)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True

    if not style_exists(doc, "Appendix Source"):
        style = styles.add_style("Appendix Source", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.italic = True
        style.font.size = Pt(9.5)
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(8)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = None
    for child in tbl_pr.iterchildren():
        if child.tag == qn("w:tblBorders"):
            borders = child
            break
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge_tag = qn(f"w:{edge}")
        edge_element = None
        for child in borders.iterchildren():
            if child.tag == edge_tag:
                edge_element = child
                break
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            borders.append(edge_element)
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:sz"), "8")
        edge_element.set(qn("w:space"), "0")
        edge_element.set(qn("w:color"), "808080")


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                shade_cell(cell, "D9E2F3")

            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    if row_idx == 0:
                        run.bold = True


def format_box_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def apply_source_text(paragraph, text):
    paragraph.text = text
    paragraph.style = "Appendix Source"


def format_score_line(paragraph):
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(4.8), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.line_spacing = 1.5


def save_document(doc):
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    candidate_paths = [OUTPUT_PATH] + [
        OUTPUT_PATH.with_name(f"{OUTPUT_PATH.stem}_v{index}{OUTPUT_PATH.suffix}")
        for index in range(2, 10)
    ]
    last_error = None
    for candidate_path in candidate_paths:
        try:
            doc.save(str(candidate_path))
            return candidate_path
        except PermissionError as exc:
            last_error = exc
            continue
    raise last_error


def enhance_appendices(doc):
    ensure_appendix_styles(doc)

    appendix_headings = [
        find_paragraph_exact(doc, "Παράρτημα Α: Συμπληρωμένο OCAI", style_name="Heading 1"),
        find_paragraph_exact(doc, "Παράρτημα Β: Συμπληρωμένο ερωτηματολόγιο οργανωσιακών τύπων του Mintzberg", style_name="Heading 1"),
    ]
    for paragraph in appendix_headings:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)

    appendix_page_titles = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith("Παράρτημα Α - Σελίδα")
        or paragraph.text.strip().startswith("Παράρτημα Β - Σελίδα")
    ]
    for paragraph in reversed(appendix_page_titles):
        delete_paragraph(paragraph)

    for text in ["Σύνοψη αποτελεσμάτων OCAI", "Αναλυτικά scores ανά διάσταση"]:
        paragraph = find_paragraph_exact(doc, text)
        paragraph.style = "Appendix Section Heading"

    for text in [
        "Απαντήσεις ερωτήσεων 13–25",
        "Απαντήσεις ερωτήσεων 26–40",
        "Σύνοψη βαθμολογιών του Mintzberg Organisational Types Audit",
        "Προφίλ αποτελεσμάτων και σύντομη ερμηνεία",
    ]:
        paragraph = find_paragraph_exact(doc, text)
        paragraph.style = "Appendix Section Heading"

    appendix_b_heading = find_paragraph_exact(
        doc,
        "Παράρτημα Β: Συμπληρωμένο ερωτηματολόγιο οργανωσιακών τύπων του Mintzberg",
        style_name="Heading 1",
    )
    appendix_b_first_existing_section = find_paragraph_exact(doc, "Απαντήσεις ερωτήσεων 13–25")
    delete_paragraphs_between(doc, appendix_b_heading, appendix_b_first_existing_section)

    intro_title = insert_paragraph_after(
        appendix_b_heading,
        "ORGANISATIONAL TYPES AUDIT",
        "Appendix Section Heading",
    )
    intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_text_1 = insert_paragraph_after(
        intro_title,
        "This questionnaire studies organizational character. You can assess a unit, site, company "
        "or non-commercial organization. Everyone filling out the questionnaire must write a brief "
        "but unmistakable definition of the organization being studied in the box below:",
        "Normal",
    )
    org_label = insert_paragraph_after(intro_text_1, "The organization being studied is:", "Normal")
    org_table = insert_table_after_paragraph(doc, org_label, 1, 1, "Normal Table")
    org_table.cell(0, 0).text = (
        "Intragen is a pan-European identity and access management (IAM) consultancy and services "
        "company delivering expert-led, project-based identity and security solutions through "
        "collaborative, cross-functional teams."
    )
    format_box_table(org_table)

    intro_text_2 = insert_paragraph_after_table(
        org_table,
        "With reference only to the organization described above, complete the questionnaire. "
        "Below are listed forty statements that could describe an organization. You must evaluate "
        "the accuracy of each statement. Allocate the points based on:",
        "Normal",
    )
    score_line_1 = insert_paragraph_after(intro_text_2, "This statement is true\t3 points", "Normal")
    score_line_2 = insert_paragraph_after(score_line_1, "This statement is partly true\t2 points", "Normal")
    score_line_3 = insert_paragraph_after(score_line_2, "This statement is untrue\t0 points", "Normal")
    for paragraph in [score_line_1, score_line_2, score_line_3]:
        format_score_line(paragraph)

    items_1_12 = [
        ("ITEMS", "POINTS"),
        ("1. The organization is directly controlled by one person", "0"),
        ("2. The organization is controlled through an elaborate hierarchy", "0"),
        ("3. Management tends to be fairly weak because the organization is ruled by largely independent professionals", "3"),
        ("4. This organization operates as a headquarters and allows operating units a good deal of freedom; provided they perform well.", "2"),
        ("5. The work requires so much creativity that 'experts' must get together to decide how things will be done.", "3"),
        ("6. The organization is always reorganizing to suit different projects", "3"),
        ("7. The organization has a number of self-contained divisions", "0"),
        ("8. There are 'professionals' at every level who make the most of their own decisions", "3"),
        ("9. There are formal rules and regulations governing almost all eventualities", "0"),
        ("10. The owner or chief executive personally makes all of the key decisions", "0"),
        ("11. There is one 'boss' who drives the whole organization to respond quickly to changes in the environment", "0"),
        ("12. Comprehensive and formal planning takes place before changes in the work organization are made", "2"),
    ]
    first_items_table = insert_table_after_paragraph(doc, score_line_3, len(items_1_12), 2, "Normal Table")
    for row_idx, row_values in enumerate(items_1_12):
        for col_idx, value in enumerate(row_values):
            first_items_table.cell(row_idx, col_idx).text = value
    format_table(first_items_table)
    first_items_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    first_items_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_items_table.columns[1].cells[0].width = first_items_table.columns[1].cells[0].width

    figure_caption = find_last_paragraph_exact(
        doc,
        "Σχήμα 2. Προφίλ οργανωσιακής δομής της εταιρείας βάσει του ερωτηματολογίου του Mintzberg",
    )
    figure_caption.text = "Σχήμα Β1. Προφίλ οργανωσιακής δομής της εταιρείας βάσει του ερωτηματολογίου του Mintzberg"
    figure_caption.style = "Appendix Caption"

    figure_source = find_paragraph_exact(
        doc,
        "Πηγή: Συμπληρωμένο ερωτηματολόγιο του Mintzberg και δημιουργία του συγγραφέα..",
    )
    apply_source_text(
        figure_source,
        "Πηγή: Συμπληρωμένο ερωτηματολόγιο οργανωσιακών τύπων του Mintzberg της παρούσας μελέτης και δημιουργία του συγγραφέα.",
    )

    figure_explanation = find_paragraph_exact(
        doc,
        "Η υψηλότερη βαθμολογία εμφανίζεται στην Επαγγελματική γραφειοκρατεία (24), με πολύ κοντινή την Καινοτομική πρασαρμογή (23). Αυτό υποδηλώνει έναν οργανισμό που βασίζεται έντονα στην εξειδικευμένη γνώση και παράλληλα χρειάζεται ευελιξία και δημιουργική προσαρμογή στα έργα.",
    )
    figure_explanation.text = (
        "The highest score appears in Professional Bureaucracy (24), with Adhocracy (23) very close "
        "behind. This indicates an organization that relies heavily on specialized knowledge while "
        "also requiring flexibility and creative adaptation in project work."
    )

    for text in [
        "Τρέχουσα κατάσταση – Μέσοι όροι προφίλ κουλτούρας",
        "Επιθυμητή κατάσταση – Μέσοι όροι προφίλ κουλτούρας",
        "Τρέχουσα Κατάσταση",
        "Επιθυμητή κατάσταση",
    ]:
        delete_paragraph(find_paragraph_exact(doc, text))

    for old, new in [
        ("Απαντήσεις ερωτήσεων 13–25", "Responses to questions 13-25"),
        ("Απαντήσεις ερωτήσεων 26–40", "Responses to questions 26-40"),
        ("Σύνοψη βαθμολογιών του Mintzberg Organisational Types Audit", "Summary of scores from the Mintzberg Organisational Types Audit"),
        ("Προφίλ αποτελεσμάτων και σύντομη ερμηνεία", "Results profile and brief interpretation"),
        ("Ο παρακάτω πίνακας δείχνει τα items που αντιστοιχούν σε κάθε οργανωσιακό τύπο και το συνολικό score.", "The table below shows the items associated with each organisational type and the total score."),
    ]:
        paragraph = find_paragraph_exact(doc, old)
        paragraph.text = new

    mintzberg_header_updates = {
        6: ["Q", "Statement", "Score"],
        7: ["Q", "Statement", "Score"],
        8: ["Code", "Organisational type", "Item scores", "Total"],
        9: ["Type", "Score", "Brief interpretation"],
    }
    for table_index, headers in mintzberg_header_updates.items():
        if table_index < len(doc.tables):
            table = doc.tables[table_index]
            for col_idx, header in enumerate(headers):
                table.cell(0, col_idx).text = header

    interpretation_rows = [
        ("Simple Structure", "0", "Very limited evidence of strongly centralized control by a single individual."),
        ("Machine Bureaucracy", "6", "Low emphasis on routine standardization and highly formalized procedures."),
        ("Professional Bureaucracy", "24", "Strong reliance on specialized expertise, autonomy, and professional standards."),
        ("Divisionalised Form", "10", "Some indications of relatively autonomous units and performance monitoring."),
        ("Adhocracy", "23", "Strong evidence of flexibility, team-based co-ordination, and innovation."),
    ]
    if 9 < len(doc.tables):
        interpretation_table = doc.tables[9]
        for row_idx, row_values in enumerate(interpretation_rows, start=1):
            for col_idx, value in enumerate(row_values):
                interpretation_table.cell(row_idx, col_idx).text = value

    appendix_a_source = "Πηγή: Συμπληρωμένο ερωτηματολόγιο OCAI της παρούσας μελέτης και επεξεργασία του συγγραφέα."
    appendix_b_source = (
        "Πηγή: Συμπληρωμένο ερωτηματολόγιο οργανωσιακών τύπων του Mintzberg της παρούσας μελέτης "
        "και επεξεργασία του συγγραφέα."
    )
    table_captions = [
        "Πίνακας Α1. Τρέχουσα κατάσταση - μέσοι όροι προφίλ κουλτούρας",
        "Πίνακας Α2. Επιθυμητή κατάσταση - μέσοι όροι προφίλ κουλτούρας",
        "Πίνακας Α3. Αναλυτικά scores ανά διάσταση - τρέχουσα κατάσταση",
        "Πίνακας Α4. Αναλυτικά scores ανά διάσταση - επιθυμητή κατάσταση",
        "Πίνακας Β1. Απαντήσεις του συμπληρωμένου ερωτηματολογίου για τις ερωτήσεις 1-12",
        "Πίνακας Β2. Απαντήσεις του συμπληρωμένου ερωτηματολογίου για τις ερωτήσεις 13-25",
        "Πίνακας Β3. Απαντήσεις του συμπληρωμένου ερωτηματολογίου για τις ερωτήσεις 26-40",
        "Πίνακας Β4. Σύνοψη βαθμολογιών ανά οργανωσιακό τύπο του Mintzberg",
        "Πίνακας Β5. Συνοπτική ερμηνεία των αποτελεσμάτων ανά οργανωσιακό τύπο",
    ]

    data_table_index = 0
    for table in doc.tables:
        if table._tbl is org_table._tbl:
            continue
        format_table(table)
        caption_paragraph = insert_paragraph_after_table(table, table_captions[data_table_index], "Appendix Caption")
        source_text = appendix_a_source if data_table_index <= 3 else appendix_b_source
        source_paragraph = insert_paragraph_after(caption_paragraph, "", "Appendix Source")
        apply_source_text(source_paragraph, source_text)
        data_table_index += 1


def main():
    doc = Document(str(INPUT_PATH))

    insertions = [
        (
            "Μια πιο πρακτική προσέγγιση δίνεται από τους Cameron και Quinn",
            [
                "Η πιο πρόσφατη βιβλιογραφία ενισχύει αυτή την προσέγγιση, καθώς δείχνει ότι η οργανωσιακή κουλτούρα επηρεάζει όχι μόνο τη συνοχή και τον συντονισμό, αλλά και τη δέσμευση, τη δημιουργικότητα και τη δυνατότητα του οργανισμού να μετατρέπει ιδέες σε πρακτική καινοτομία [Ding and Hong, 2025; Clouet et al., 2024]. Παράλληλα, η βιβλιογραφική χαρτογράφηση της κουλτούρας καινοτομίας δείχνει ότι οι οργανισμοί τείνουν να είναι πιο προσαρμοστικοί όταν συνδυάζουν την έμφαση στις σχέσεις με χώρο για πρωτοβουλία, πειραματισμό και συνεργασία [Pineda-Celaya et al., 2023].",
            ],
        ),
        (
            "Ενδιαφέρον φαίνεται να έχει και η εστίαση  στην προσωπική ανάπτυξη και στη διαρκή μάθηση.",
            [
                "Στην ίδια λογική, πρόσφατες συστηματικές ανασκοπήσεις δείχνουν ότι η μάθηση υποστηρίζεται περισσότερο όταν ο οργανισμός δημιουργεί σταθερούς μηχανισμούς ανταλλαγής γνώσης, κοινότητες μάθησης και πρακτικές ανάπτυξης δεξιοτήτων που επαναλαμβάνονται μέσα στον χρόνο, αντί να βασίζεται μόνο σε αποσπασματικές εκπαιδευτικές δράσεις [Zamiri and Esmaeili, 2024a; Zamiri and Esmaeili, 2024b]. Επιπλέον, σε περιβάλλοντα υπηρεσιών, η αξιοποίηση της άρρητης γνώσης και της οργανωσιακής μάθησης συνδέεται θετικά με τη συνολική απόδοση, κάτι που καθιστά ιδιαίτερα σημαντική την καθημερινή μεταφορά εμπειρίας από πιο έμπειρους προς νεότερους εργαζομένους [AlMulhim, 2020].",
            ],
        ),
        (
            "Η μελέτη δεν περιορίστηκε μόνο στα αποτελέσματα των ερωτηματολογίων.",
            [
                "Με αυτή την έννοια, τα δύο συμπληρωμένα ερωτηματολόγια αντιμετωπίστηκαν και ως πρωτογενές υλικό της παρούσας εργασίας, καθώς δεν χρησιμοποιήθηκαν μόνο για την παραγωγή διαγραμμάτων, αλλά και ως δομημένο σημείο αναφοράς για την ερμηνεία της οργανωσιακής δομής και της κουλτούρας της εταιρείας. Τα πλήρη ερωτηματολόγια παρατίθενται στα Παραρτήματα Α και Β.",
            ],
        ),
        (
            "Παράλληλα, φαίνεται να δίνεται σημασία και στη φυσική παρουσία στο γραφείο κατά το πρώτο διάστημα.",
            [
                "Η έμφαση στην άτυπη υποστήριξη, στην εύκολη πρόσβαση σε καθοδήγηση και στη στενότερη καθημερινή επαφή κατά το πρώτο στάδιο ένταξης είναι συμβατή με νεότερα ευρήματα που συνδέουν τις συμμετοχικές μορφές ηγεσίας και τις υψηλής εμπλοκής πρακτικές ανθρώπινου δυναμικού με καλύτερη απόδοση και πιο ενεργή συμβολή των εργαζομένων στην καθημερινή λειτουργία [AlMulhim and Mohammed, 2023; Cahyadi et al., 2022].",
            ],
        ),
        (
            "Η διάσταση της ανάπτυξης ενισχύεται και από τον ρόλο του επικεφαλή ομάδας στις μηνιαίες ατομικές συναντήσεις.",
            [
                "Η εικόνα αυτή ταιριάζει και με τη λογική της αναπτυξιακής ηγεσίας, σύμφωνα με την οποία η καθοδήγηση είναι πιο αποτελεσματική όταν βοηθά τον εργαζόμενο να βλέπει πιο καθαρά την επόμενη φάση ανάπτυξής του, να μαθαίνει μέσα από σταδιακές προκλήσεις και να ενισχύει την ανθεκτικότητά του μέσα στην εργασία [Franken et al., 2021; Franken et al., 2023].",
            ],
        ),
        (
            "Στην ίδια κατεύθυνση κινείται και η πρόθεση της εταιρείας να επενδύσει πιο συστηματικά στη γνώση γύρω από την τεχνητή νοημοσύνη μέσα από μια Ακαδημία Τεχνητής Νοημοσύνης.",
            [
                "Παράλληλα, η σύνδεση της εκπαίδευσης με κοινές πρακτικές ανταλλαγής γνώσης είναι ιδιαίτερα σημαντική, καθώς η πρόσφατη έρευνα δείχνει ότι τα εργαστήρια, οι μικρές κοινότητες μάθησης και τα επαναλαμβανόμενα σχήματα διαμοιρασμού εμπειρίας διευκολύνουν τη διατήρηση και τη μεταφορά κρίσιμης τεχνογνωσίας [Zamiri and Esmaeili, 2024a; Zamiri and Esmaeili, 2024b].",
            ],
        ),
        (
            "Ανάλογη εικόνα προκύπτει και από τα ευρήματα του OCAI.",
            [
                "Η νεότερη βιβλιογραφία ενισχύει την παραπάνω εικόνα, καθώς δείχνει ότι η καινοτομική κουλτούρα λειτουργεί πιο αποτελεσματικά όταν συνδυάζει συνεργασία, εμπιστοσύνη και σαφή προσανατολισμό στη δημιουργία νέων λύσεων. Σε τέτοια περιβάλλοντα, η δέσμευση των εργαζομένων και η καινοτομική τους συμβολή φαίνεται να ενισχύονται αμοιβαία [Ding and Hong, 2025; Clouet et al., 2024; Pineda-Celaya et al., 2023].",
            ],
        ),
        (
            "Η κατεύθυνση αυτή συνδέεται και με τη σύγχρονη συζήτηση γύρω από την καινοτομία και την αξιοποίηση της τεχνητής νοημοσύνης στους οργανισμούς.",
            [
                "Σε επίπεδο καθημερινής συμπεριφοράς, η καινοτομία ενισχύεται περισσότερο όταν η ηγεσία καλλιεργεί αίσθηση εμπιστοσύνης, ενδυναμώνει την αυτοπεποίθηση των εργαζομένων και δημιουργεί χώρο για ενεργή εμπλοκή στην επίλυση προβλημάτων. Αυτό έχει συνδεθεί τόσο με την ηθική όσο και με τη συμπεριληπτική ηγεσία [Uppathampracha and Liu, 2022; Bunkaewsuk et al., 2024; AlMulhim and Mohammed, 2023].",
            ],
        ),
        (
            "Σε αυτό το πλαίσιο, φαίνεται ότι διαθέτει αρκετά στοιχεία που ενισχύουν τα μη χρηματικά κίνητρα.",
            [
                "Η εικόνα αυτή συμφωνεί και με πρόσφατες μελέτες που δείχνουν ότι η δέσμευση και η ικανοποίηση των εργαζομένων ενισχύονται όταν το περιβάλλον εργασίας συνδέεται με αίσθηση υποστήριξης, σχέσεις εμπιστοσύνης και πραγματικές δυνατότητες εξέλιξης, και όχι μόνο με τυπικές παροχές [Sypniewska et al., 2023; Ding and Hong, 2025].",
            ],
        ),
        (
            "Εξίσου σημαντικός φαίνεται να είναι και ο τρόπος ανατροφοδότησης και αξιολόγησης.",
            [
                "Η πρόσφατη βιβλιογραφία για την αξιολόγηση της απόδοσης κινείται επίσης προς αυτή την κατεύθυνση. Πιο συγκεκριμένα, δείχνει ότι η δικαιοσύνη του συστήματος αξιολόγησης και ο αναπτυξιακός του σκοπός επηρεάζουν θετικά την εργασιακή δέσμευση, ενώ οι πιο συχνές και ουσιαστικές συζητήσεις ανατροφοδότησης καθιστούν τη διαχείριση της απόδοσης πιο χρήσιμη για τον εργαζόμενο [Vidé et al., 2023; Micacchi et al., 2024; Barbieri et al., 2023; O'Kane et al., 2023].",
            ],
        ),
        (
            "Σημαντικό ρόλο στα κίνητρα φαίνεται να παίζει και το οικονομικό σκέλος.",
            [
                "Αντίστοιχα, η σχέση ανάμεσα στην οργανωσιακή κουλτούρα, στα κίνητρα και στην απόδοση φαίνεται να γίνεται ισχυρότερη όταν μεσολαβεί η ικανοποίηση από την εργασία, κάτι που δείχνει ότι η θετική εμπειρία του εργαζομένου δεν αποτελεί απλώς δευτερεύον στοιχείο, αλλά βασικό μηχανισμό οργανωσιακής αποτελεσματικότητας [Al-Ansi et al., 2023].",
            ],
        ),
        (
            "Επιπλέον, αυτή η κίνηση μπορεί να ενισχύσει και την ομαδικότητα.",
            [
                "Μια τέτοια ομάδα θα μπορούσε να λειτουργήσει και ως μικρός οργανωσιακός κόμβος καινοτομίας. Η σχετική βιβλιογραφία δείχνει ότι η καινοτομία σταθεροποιείται περισσότερο όταν συνδέεται με ανοιχτή συνεργασία, διαλειτουργική συμμετοχή και πρακτικές που επιτρέπουν στους εργαζομένους να μετατρέπουν ιδέες σε βελτιώσεις της καθημερινής εργασίας [Clouet et al., 2024; Pineda-Celaya et al., 2023; Uppathampracha and Liu, 2022].",
            ],
        ),
        (
            "Στην ίδια κατεύθυνση, θα ήταν χρήσιμο να επανενεργοποιηθούν πιο συστηματικά τα εργαστήρια ανά δύο εβδομάδες και τα βίντεο ανταλλαγής γνώσης,",
            [
                "Η κατεύθυνση αυτή υποστηρίζεται και από νεότερες μελέτες που αναδεικνύουν τη σημασία των κοινοτήτων μάθησης, των σταθερών workshop και των οργανωμένων πρακτικών ανταλλαγής γνώσης για τη μετάβαση από την εκπαίδευση στην πράξη [Zamiri and Esmaeili, 2024a; Zamiri and Esmaeili, 2024b]. Παράλληλα, η συστηματική αξιοποίηση της εμπειρικής γνώσης μέσα στην καθημερινή εργασία φαίνεται να είναι ιδιαίτερα σημαντική σε περιβάλλοντα υπηρεσιών και τεχνογνωσίας [AlMulhim, 2020].",
            ],
        ),
        (
            "Η παρέμβαση αυτή είναι σημαντική, γιατί η πρακτική εμπειρία δεν συμβάλλει μόνο στην τεχνική ανάπτυξη,",
            [
                "Η σταδιακή ανάληψη ευθύνης είναι επίσης συμβατή με τη λογική της αναπτυξιακής ηγεσίας, καθώς δίνει στους νεότερους εργαζομένους τη δυνατότητα να ενισχύουν την αυτοπεποίθηση και την ανθεκτικότητά τους χωρίς απότομη έκθεση σε υψηλές απαιτήσεις [Franken et al., 2021; Franken et al., 2023].",
            ],
        ),
        (
            "Παράλληλα, η εξαμηνιαία αξιολόγηση θα μπορούσε να συνδεθεί πιο άμεσα με αυτό το πλάνο.",
            [
                "Η προσέγγιση αυτή είναι κοντά και στη σύγχρονη λογική της διαχείρισης της απόδοσης, η οποία αντιμετωπίζει την αξιολόγηση περισσότερο ως συνεχή αναπτυξιακή συζήτηση και λιγότερο ως μεμονωμένο διοικητικό έλεγχο [O'Kane et al., 2023; Vidé et al., 2023].",
            ],
        ),
        (
            "Μια ακόμη πρακτική βελτίωση θα ήταν να υπάρχει πιο σταθερή δομή και στην ανατροφοδότηση από τους άλλους εργαζομένους.",
            [
                "Επιπλέον, όταν η διαδικασία γίνεται αντιληπτή ως δίκαιη, σαφής και χρήσιμη για την εξέλιξη του εργαζομένου, είναι πιο πιθανό να ενισχύει τη δέσμευση και την αποδοχή του συστήματος αξιολόγησης [Barbieri et al., 2023; Micacchi et al., 2024].",
            ],
        ),
        (
            "Με βάση τα παραπάνω, οι προτάσεις βελτίωσης εστίασαν στην ανάπτυξη εσωτερικών εργαλείων παραγωγικότητας,",
            [
                "Η πρόσφατη βιβλιογραφία ενισχύει αυτή την κατεύθυνση, δείχνοντας ότι η ανάπτυξη, η καινοτομία και η δέσμευση αλληλοενισχύονται όταν ο οργανισμός συνδέει υποστηρικτική κουλτούρα, αναπτυξιακή ηγεσία και δίκαιη ανατροφοδότηση [Ding and Hong, 2025; Franken et al., 2023; Micacchi et al., 2024].",
            ],
        ),
    ]

    for prefix, texts in insertions:
        insert_after_prefix(doc, prefix, texts)

    bibliography_entries = [
        "Al-Ansi, A.M., Jaboob, M. and Awain, A.M.S.B. (2023) Examining the mediating role of job satisfaction between motivation, organizational culture, and employee performance in higher education: a case study in the Arab region. Education Science and Management, 1(1), pp. 30-42.",
        "AlMulhim, A.F. (2020) The effect of tacit knowledge and organizational learning on financial performance in service industry. Management Science Letters, 10(10), pp. 2211-2220.",
        "AlMulhim, A.F. and Mohammed, S.M. (2023) The impact of inclusive leadership on innovative work behavior: a mediated moderation model. Leadership & Organization Development Journal, 44(7), pp. 907-926.",
        "Barbieri, M., Micacchi, L., Vide, F. and Valotti, G. (2023) The performance of performance appraisal systems: a theoretical framework for public organizations. Review of Public Personnel Administration, 43(1), pp. 104-129.",
        "Buchanan, D.A. and Huczynski, A.A. (2013) Organizational Behaviour. 8th edn. Harlow: Pearson.",
        "Cahyadi, A., Marwa, T., Hágen, I., Siraj, M.N., Santati, P., Poór, J. and Szabó, K. (2022) Leadership styles, high-involvement human resource management practices, and individual employee performance in small and medium enterprises in the digital era. Economies, 10(7), article no. 162.",
        "Cameron, K.S. and Quinn, R.E. (2006) Diagnosing and Changing Organizational Culture: Based on the Competing Values Framework. Revised edn. San Francisco, CA: Jossey-Bass.",
        "Clouet, M.E., Alfaro-Tanco, J.A. and Recalde, M. (2024) Building a framework to promote corporate social innovation: a view from organizational culture. Revista Empresa y Humanismo, 27(2), pp. 9-49.",
        "Ding, J. and Hong, G. (2025) Fostering loyalty and creativity: how organizational culture shapes employee commitment and innovation in South Korean firms. Behavioral Sciences, 15(4), article no. 529.",
        "Franken, E., Plimmer, G., Malinen, S.K., Bryson, J. and Berman, E.M. (2021) Building people up: growth-oriented leadership in the public sector. Australian Journal of Public Administration, 80(4), pp. 661-689.",
        "Franken, E., Plimmer, G. and Malinen, S. (2023) Growth-oriented management and employee outcomes: employee resilience as a mechanism for growth. Leadership & Organization Development Journal, 44(5), pp. 627-642.",
        "Mintzberg, H. (1979) The Structuring of Organizations: A Synthesis of the Research. Englewood Cliffs, NJ: Prentice-Hall.",
        "Mintzberg Organisational Types Audit questionnaire (n.d.) Unpublished questionnaire completed for the purposes of the present study; see Appendix B.",
        "Micacchi, L., Vide, F., Giacomelli, G. and Barbieri, M. (2024) Performance appraisal justice and employees' work engagement in the public sector: making the most of performance appraisal design. Public Administration, 102(3), pp. 815-840.",
        "O'Kane, P., McCracken, M. and Brown, T. (2023) Reframing the performance management system: a conversations perspective. Personnel Review, 52(5), pp. 1579-1596.",
        "Organizational Culture Assessment Instrument (OCAI) questionnaire (n.d.) Unpublished questionnaire completed for the purposes of the present study; see Appendix A.",
        "Pineda-Celaya, L., Andrés-Reina, M.P. and González-Pérez, M. (2023) Detection of the lines of research in favor of the implementation and development of organizational culture of innovation through a bibliometric analysis. Innovar, 33(89), pp. 161-176.",
        "Schein, E.H. (2010) Organizational Culture and Leadership. 4th edn. San Francisco, CA: Jossey-Bass.",
        "Senge, P.M., Kleiner, A., Roberts, C., Ross, R.B. and Smith, B.J. (1994) The Fifth Discipline Fieldbook: Strategies and Tools for Building a Learning Organization. London: Nicholas Brealey Publishing.",
        "Sułkowski, Ł. and Lenart-Gansiniec, R. (2025) Paradigm Shifts in Management Theory. 1st edn. Abingdon and New York: Routledge.",
        "Sypniewska, B., Baran, M. and Kłos, M. (2023) Work engagement and employee satisfaction in the practice of sustainable human resource management - based on the study of Polish employees. International Entrepreneurship and Management Journal, 19(3), pp. 1069-1100.",
        "Uppathampracha, R. and Liu, G. (2022) Leading for innovation: self-efficacy and work engagement as sequential mediation relating ethical leadership and innovative work behavior. Behavioral Sciences, 12(8), article no. 266.",
        "Vide, F., Micacchi, L., Barbieri, M. and Valotti, G. (2023) The renaissance of performance appraisal: engaging public employees through perceived developmental purpose and justice. Review of Public Personnel Administration, 43(4), pp. 623-651.",
        "Zamiri, M. and Esmaeili, A. (2024a) Methods and technologies for supporting knowledge sharing within learning communities: a systematic literature review. Administrative Sciences, 14(1), article no. 17.",
        "Zamiri, M. and Esmaeili, A. (2024b) Strategies, methods, and supports for developing skills within learning communities: a systematic review of the literature. Administrative Sciences, 14(9), article no. 231.",
    ]

    replace_bibliography(doc, bibliography_entries)

    replacements = [
        ("αντιδρασης στις αλλαγές των μελλών τους", "αντίδρασης στις αλλαγές των μελών τους"),
        ("μέσα στον οργανισμο [", "μέσα στον οργανισμό ["),
        ("να κατανοήσουμε πως του ένας οργανισμός", "να κατανοήσουμε πώς ένας οργανισμός"),
        ("το κοινο όραμα", "το κοινό όραμα"),
        ("του τρόπου  διαμόρφωσης της συνεργασίας, της καθοδήγησης, και της ένταξης νέων εργαζομένων και η διαχείριση της απόδοσης.", "του τρόπου διαμόρφωσης της συνεργασίας, της καθοδήγησης, της ένταξης νέων εργαζομένων και της διαχείρισης της απόδοσης."),
        ("ιδιαίτερα  όσων αφορά", "ιδιαίτερα όσον αφορά"),
        ("σε έναν επικεφαλή ομάδας", "σε έναν επικεφαλής ομάδας"),
        ("του επικεφαλή ομάδας", "του επικεφαλής ομάδας"),
        ("6.3 Ενίσχυση του αναπτυξιακού ρόλου του επικεφαλή ομάδας και καλύτερη αξιοποίηση της αξιολόγησης", "6.3 Ενίσχυση του αναπτυξιακού ρόλου του επικεφαλής ομάδας και καλύτερη αξιοποίηση της αξιολόγησης"),
        ("θα αναλυθούν πιο αναλυτικά στην επόμενη ενότητα.", "θα αναλυθούν εκτενέστερα στην επόμενη ενότητα."),
        ("καθώς συνδέονται άμεσα με την ικανότητα της να διατηρεί να εξελίσει και να αξιοποιεί αποτελεσματικά το ανθρώπινο δυναμικό της.", "καθώς συνδέονται άμεσα με την ικανότητά της να διατηρεί, να εξελίσσει και να αξιοποιεί αποτελεσματικά το ανθρώπινο δυναμικό της."),
        ("Παρόλα αυτά", "Παρ' όλα αυτά"),
        ("[Uppathampracha and Liu, 2022; Bunkaewsuk et al., 2024; AlMulhim and Mohammed, 2023].", "[Uppathampracha and Liu, 2022; AlMulhim and Mohammed, 2023]."),
        ("[Vidé et al., 2023; Micacchi et al., 2024; Barbieri et al., 2023; O'Kane et al., 2023].", "[Vide et al., 2023; Micacchi et al., 2024; Barbieri et al., 2023; O'Kane et al., 2023]."),
        ("[O'Kane et al., 2023; Vidé et al., 2023].", "[O'Kane et al., 2023; Vide et al., 2023]."),
    ]
    replace_text_in_paragraphs(doc, replacements)

    paragraph_updates = [
        (
            "Η εργασία εστιάζει στην εξέταση της λειτουργίας της εταιρείας",
            "Η εργασία εξετάζει τη λειτουργία της εταιρείας μέσα από την οργανωσιακή δομή και την οργανωσιακή κουλτούρα, με έμφαση στο πώς υποστηρίζονται η εξέλιξη των εργαζομένων, η προσαρμογή στην αλλαγή και το εργασιακό κλίμα. Για τον σκοπό αυτό αξιοποιούνται το ερωτηματολόγιο του Mintzberg και το εργαλείο των Cameron και Quinn, ώστε η ανάλυση να στηριχθεί σε πιο συστηματικά ευρήματα.",
        ),
        (
            "Η οργανωσιακή δομή αφορά τον τρόπο που ένας οργανισμός μοιράζει",
            "Η οργανωσιακή δομή αφορά τον τρόπο με τον οποίο κατανέμεται η εργασία και συντονίζονται οι ρόλοι, ώστε ο οργανισμός να λειτουργεί αποτελεσματικά. Για τον Mintzberg, δεν είναι μόνο θέμα ιεραρχίας, αλλά και του τρόπου με τον οποίο η εργασία συνδέεται στην πράξη (Mintzberg, 1979, σελ. 2-3).",
        ),
        (
            "Ο Mintzberg υποστηρίζει ότι κάθε οργανισμός πρέπει να αντιμετωπίσει",
            "Ο Mintzberg υποστηρίζει ότι κάθε οργανισμός καλείται να λύσει δύο ζητήματα: πώς θα διαιρέσει τη συνολική εργασία και πώς θα συντονίσει τα επιμέρους καθήκοντα. Για τον σκοπό αυτό περιγράφει πέντε βασικούς μηχανισμούς συντονισμού: αμοιβαία προσαρμογή, άμεση επίβλεψη, τυποποίηση διαδικασιών, τυποποίηση αποτελεσμάτων και τυποποίηση δεξιοτήτων (Mintzberg, 1979, σελ. 3-10).",
        ),
        (
            "Για τη συγκεκριμένη ανάλυση, δύο μορφές οργάνωσης του Mintzberg",
            "Για τη συγκεκριμένη ανάλυση έχουν μεγαλύτερο ενδιαφέρον η επαγγελματική γραφειοκρατία και η πιο ευέλικτη, καινοτομική μορφή οργάνωσης. Η πρώτη στηρίζεται στην εξειδίκευση και στην αυτονομία εκπαιδευμένων επαγγελματιών (Mintzberg, 1979, σελ. 349-352). Γι' αυτό η εκπαίδευση και η ανάπτυξη δεξιοτήτων αποτελούν μέρος της ίδιας της οργανωσιακής δομής (Mintzberg, 1979, σελ. 9-10, 349-366).",
        ),
        (
            "Κατά τον Schein, η κουλτούρα μπορεί να εξεταστεί",
            "Κατά τον Schein, η κουλτούρα μπορεί να εξεταστεί σε τρία επίπεδα: ορατά στοιχεία, δηλωμένες αξίες και βασικές υποκείμενες παραδοχές (Schein, 2010, σελ. 23-33). Η διάκριση αυτή είναι χρήσιμη, γιατί δείχνει ότι η κουλτούρα δεν είναι επιφανειακό χαρακτηριστικό αλλά μηχανισμός που επηρεάζει συνοχή, σταθερότητα και μάθηση (Schein, 2010, σελ. 31-33).",
        ),
        (
            "Μια πιο πρακτική προσέγγιση δίνεται από τους Cameron και Quinn",
            "Μια πιο πρακτική προσέγγιση προσφέρουν οι Cameron και Quinn μέσω του Πλαισίου Ανταγωνιστικών Αξιών, το οποίο συνδυάζει τη διάκριση ευελιξίας-ελέγχου με τη διάκριση εσωτερικής-εξωτερικής εστίασης (Cameron and Quinn, 2006, σελ. 33-36). Από αυτό προκύπτουν τέσσερις τύποι κουλτούρας, με διαφορετικές αξίες και πρότυπα ηγεσίας (Cameron and Quinn, 2006, σελ. 35-46).",
        ),
        (
            "Η παρούσα εργασία προσεγγίζει την Intragen ως μια συγκεκριμένη οργανωσιακή περίπτωση",
            "Η παρούσα εργασία προσεγγίζει την Intragen ως συγκεκριμένη οργανωσιακή περίπτωση, με στόχο να αποτυπώσει βασικά στοιχεία της λειτουργίας της και να τα ερμηνεύσει μέσα από επιλεγμένα θεωρητικά εργαλεία. Η έμφαση δίνεται σε πλευρές της καθημερινής λειτουργίας που επηρεάζουν άμεσα την εμπειρία των εργαζομένων και τη συνολική αποτελεσματικότητα.",
        ),
        (
            "Η μελέτη δεν περιορίστηκε μόνο στα αποτελέσματα των ερωτηματολογίων.",
            "Η μελέτη δεν περιορίστηκε στα αποτελέσματα των ερωτηματολογίων. Τα ευρήματα συνδέθηκαν με στοιχεία για την κατανομή του προσωπικού, την εκπαίδευση, τον ρόλο του μέντορα και του επικεφαλής ομάδας, τις πρακτικές ανατροφοδότησης και πρωτοβουλίες που σχετίζονται με μάθηση, συνεργασία και καινοτομία, καθώς και με εσωτερική πληροφορία που παρασχέθηκε από εργαζόμενο σχετικά με καθημερινές πρακτικές και οργανωσιακές διαδικασίες. Έτσι, τα ερωτηματολόγια λειτούργησαν ως αφετηρία ευρύτερης οργανωσιακής ανάγνωσης και όχι ως αυτάρκες αποτέλεσμα.",
        ),
        (
            "Με αυτή την έννοια, τα δύο συμπληρωμένα ερωτηματολόγια αντιμετωπίστηκαν",
            "Τα δύο συμπληρωμένα ερωτηματολόγια αντιμετωπίστηκαν ως πρωτογενές υλικό και παρατίθενται στα Παραρτήματα Α και Β. Επειδή, όμως, συμπληρώθηκαν από ένα στέλεχος του τεχνικού τμήματος, αποτυπώνουν κυρίως μια τεκμηριωμένη εσωτερική οπτική και όχι πλήρη οργανωσιακή συναίνεση. Για τον λόγο αυτό αξιοποιούνται ως ισχυρή ένδειξη τάσεων και αντιφάσεων, όχι ως μοναδική απόδειξη.",
        ),
        (
            "Η ανάλυση εστιάζει κυρίως στο τεχνικό τμήμα",
            "Μεθοδολογικά, μεγαλύτερο βάρος δίνεται στη λειτουργία υλοποίησης έργων, στην ένταξη νέων εργαζομένων και στην καθημερινή συνεργασία, επειδή εκεί γίνονται πιο ορατές οι σχέσεις ανάμεσα σε δομή, μάθηση και απόδοση.",
        ),
        (
            "Η οργανωσιακή εικόνα της εταιρείας διαμορφώνεται",
            "Η οργανωσιακή εικόνα της εταιρείας διαμορφώνεται κυρίως από τον τρόπο οργάνωσης της υλοποίησης έργων και από τη σύνδεσή της με τις υπόλοιπες λειτουργίες. Με βάση τα διαθέσιμα στοιχεία, η Intragen στηρίζεται σε πυρήνα εξειδικευμένης τεχνογνωσίας, γύρω από τον οποίο οργανώνονται διοικητικές, υποστηρικτικές και εμπορικές λειτουργίες.",
        ),
        (
            "Η διάρθρωση αυτή δείχνει έναν οργανισμό που στηρίζεται",
            "Η διάρθρωση αυτή δείχνει οργανισμό που στηρίζεται σε σημαντικό βαθμό στην εξειδικευμένη γνώση και στην αποτελεσματική συνεργασία μεταξύ διαφορετικών ρόλων. Γι' αυτό, πέρα από τη δομή, ιδιαίτερη σημασία αποκτούν η καθοδήγηση, η ανατροφοδότηση και η ποιότητα της καθημερινής συνεννόησης.",
        ),
        (
            "Συνολικά, η εταιρεία εμφανίζει μια οργανωσιακή διάρθρωση",
            "Το Σχήμα 1 αποτυπώνει συνοπτικά αυτή τη βασική κατανομή ρόλων και ανθρώπινου δυναμικού.",
        ),
        (
            "Η κατανομή αυτή δείχνει ότι το τεχνικό τμήμα αποτελεί",
            "Η κατανομή αυτή βοηθά να εξηγηθεί γιατί ζητήματα εκπαίδευσης, συντονισμού και μεταφοράς γνώσης επηρεάζουν άμεσα τη συνολική αποτελεσματικότητα της εταιρείας.",
        ),
        (
            "Η συνολική εικόνα της Intragen δείχνει έναν οργανισμό",
            "Η συνολική εικόνα της Intragen δείχνει οργανισμό με διακριτές λειτουργικές περιοχές και σχετικά σαφή κατανομή ρόλων. Δεν πρόκειται για απλή, αυστηρά ιεραρχική δομή, αλλά για πιο σύνθετη διάταξη στην οποία συνυπάρχουν διοικητικές, υποστηρικτικές, εμπορικές και τεχνικές λειτουργίες.",
        ),
        (
            "Η εικόνα αυτή συνδέεται και με τα ευρήματα του ερωτηματολογίου οργανωσιακών τύπων",
            "Το βασικό εύρημα του Mintzberg είναι ότι η Intragen δεν εντάσσεται καθαρά σε έναν μόνο τύπο. Η τεχνική εξειδίκευση την φέρνει κοντά στην επαγγελματική γραφειοκρατία, αλλά η εργασία σε έργα και η προσαρμογή σε διαφορετικά περιβάλλοντα πελατών την ωθούν προς πιο ευέλικτη μορφή. Αυτό είναι πλεονέκτημα για την ποιότητα και την εξυπηρέτηση πελατών, αλλά ταυτόχρονα αυξάνει την ανάγκη για ταχύτερη μάθηση και καλύτερο συντονισμό.",
        ),
        (
            "Η εικόνα αυτή συνδέεται και με τα ευρήματα του OCAI. Πιο συγκεκριμένα, η Intragen",
            "Το OCAI συμπληρώνει αυτή την εικόνα, καθώς δείχνει συνδυασμό συνεργατικής, ανθρωποκεντρικής και πιο ευέλικτης κουλτούρας. Το κρίσιμο σημείο δεν είναι απλώς η συνύπαρξη θετικών χαρακτηριστικών, αλλά η ανάγκη να διατηρηθεί η συνοχή της εταιρείας όσο μεγαλώνει χωρίς να χαθεί η προσαρμοστικότητά της.",
        ),
        (
            "Το διάγραμμα που ακολουθεί συνοψίζει τα αποτελέσματα του ερωτηματολογίου οργανωσιακών τύπων",
            "Το διάγραμμα που ακολουθεί συνοψίζει τα αποτελέσματα του ερωτηματολογίου του Mintzberg και δείχνει ότι η Intragen συνδυάζει κυρίως χαρακτηριστικά επαγγελματικής γραφειοκρατίας με ορισμένα στοιχεία μεγαλύτερης ευελιξίας. Η εικόνα αυτή αποτελεί τη βάση για την ανάλυση που ακολουθεί.",
        ),
        (
            "Η μάθηση και η ανάπτυξη αποτελούν τα πιο βασικά στοιχεία της λειτουργίας της Intragen",
            "Η μάθηση και η ανάπτυξη αποτελούν κεντρικό στοιχείο της λειτουργίας της Intragen. Η εξάμηνη περίοδος εκπαίδευσης δείχνει ότι η εταιρεία αντιμετωπίζει την ένταξη νέων εργαζομένων ως δομημένη διαδικασία και όχι ως σύντομη προσαρμογή.",
        ),
        (
            "Η υποστήριξη των νεοεισερχόμενων δεν περιορίζεται",
            "Η υποστήριξη των νεοεισερχόμενων δεν περιορίζεται στην επίσημη εκπαίδευση. Ο μέντορας καλύπτει την καθημερινή τεχνική καθοδήγηση, ενώ ο επικεφαλής ομάδας προσφέρει εποπτεία και διοικητική υποστήριξη. Αυτό δείχνει ότι η μάθηση οργανώνεται ταυτόχρονα σε τεχνικό και οργανωσιακό επίπεδο.",
        ),
        (
            "Η διάσταση της ανάπτυξης ενισχύεται και από τον ρόλο του επικεφαλής ομάδας",
            "Η διάσταση της ανάπτυξης ενισχύεται και από τις μηνιαίες ατομικές συναντήσεις με τον επικεφαλής ομάδας. Εκεί η πρόοδος συνδέεται όχι μόνο με την τρέχουσα εργασία, αλλά και με επόμενους στόχους και βήματα εξέλιξης, άρα η ανάπτυξη δεν περιορίζεται στην τεχνική κατάρτιση.",
        ),
        (
            "Η εικόνα αυτή ταιριάζει και με τη λογική της αναπτυξιακής ηγεσίας",
            "Η εικόνα αυτή συνδέεται με τη λογική της αναπτυξιακής ηγεσίας, όπου η καθοδήγηση ενισχύει μάθηση, ανθεκτικότητα και σταδιακή πρόοδο του εργαζομένου (Franken et al., 2021; Franken et al., 2023).",
        ),
        (
            "Η εικόνα αυτή είναι συνεπής και με τα ευρήματα του ερωτηματολογίου του Mintzberg",
            "Το εύρημα του Mintzberg ενισχύει αυτή την εικόνα, αλλά αναδεικνύει και μια βασική εξάρτηση: σε οργανισμούς που στηρίζονται στην εξειδίκευση, η ποιότητα του αποτελέσματος εξαρτάται από το πόσο γρήγορα οι νέοι εργαζόμενοι μετατρέπουν την εκπαίδευση σε παραγωγική συμμετοχή. Άρα, το ισχυρό σημείο της Intragen είναι η ανάπτυξη δεξιοτήτων, αλλά το ίδιο στοιχείο γίνεται πίεση αν η μετάβαση στην πράξη καθυστερεί.",
        ),
        (
            "Παράλληλα, η μάθηση στην Intragen δεν φαίνεται",
            "Παράλληλα, η μάθηση δεν περιορίζεται στην τυπική εκπαίδευση. Η αυξημένη φυσική παρουσία στο γραφείο στην αρχή διευκολύνει άτυπη μάθηση, γρήγορη υποστήριξη και σταδιακή έκθεση σε πραγματικές συνθήκες εργασίας.",
        ),
        (
            "Στην ίδια κατεύθυνση κινείται και η πρόθεση της εταιρείας να επενδύσει",
            "Η πρόθεση για Ακαδημία Τεχνητής Νοημοσύνης δείχνει προσανατολισμό στις μελλοντικές ανάγκες του οργανισμού και σύνδεση της μάθησης με την παραγωγικότητα. Ωστόσο, ως πρόθεση και όχι ακόμη εδραιωμένη πρακτική, αποτελεί περισσότερο αναπτυξιακή δυνατότητα παρά ήδη ώριμη οργανωσιακή ικανότητα.",
        ),
        (
            "Η παραπάνω εικόνα συνδέεται και με την προσέγγιση του Senge",
            "Αυτό ταιριάζει και με τον Senge, για τον οποίο η μάθηση αφορά τη συνολική ικανότητα του οργανισμού να αναπτύσσει ανθρώπους, να ενισχύει τη συνεργασία και να βελτιώνεται διαρκώς (Senge et al., 1994).",
        ),
        (
            "Συνολικά, η Intragen φαίνεται να διαθέτει ένα οργανωμένο πλαίσιο μάθησης",
            "Συνολικά, η Intragen διαθέτει ισχυρή μαθησιακή υποδομή. Το κρίσιμο εύρημα είναι ότι το αδύναμο σημείο δεν βρίσκεται στην αρχική εκπαίδευση, αλλά στη γρήγορη μετατροπή της σε ουσιαστική συμμετοχή σε έργα και σε νωρίτερη επαγγελματική αυτονομία.",
        ),
        (
            "Η προσαρμοστικότητα και η καινοτομία είναι επίσης σημαντικές",
            "Η προσαρμοστικότητα και η καινοτομία είναι κρίσιμες για την Intragen, επειδή δραστηριοποιείται σε απαιτητικό τεχνολογικό χώρο όπου τα έργα και οι ανάγκες πελατών δεν ακολουθούν πλήρως τυποποιημένο μοντέλο.",
        ),
        (
            "Η εικόνα που προκύπτει είναι ότι ο οργανισμός",
            "Η εταιρεία δεν λειτουργεί μόνο με όρους σταθερότητας και διαδικασιών, αλλά χρειάζεται ουσιαστική ευελιξία. Η εργασία οργανώνεται γύρω από έργα και απαιτεί συνεργασία, προσαρμογή και επίλυση προβλημάτων σε διαφορετικά περιβάλλοντα πελατών.",
        ),
        (
            "Η παραπάνω εικόνα συνδέεται με τα ευρήματα του ερωτηματολογίου του Mintzberg",
            "Το βασικό εύρημα του Mintzberg είναι ότι η Intragen συνδυάζει επαγγελματική γραφειοκρατία με στοιχεία ευέλικτης και καινοτομικής οργάνωσης. Αυτό δίνει τεχνική ποιότητα και ευελιξία ταυτόχρονα, αλλά σημαίνει ότι ο οργανισμός δεν μπορεί να στηριχθεί μόνο σε τυπικές διαδικασίες και χρειάζεται γρήγορο συντονισμό και μάθηση στην πράξη.",
        ),
        (
            "Ανάλογη εικόνα προκύπτει και από τα ευρήματα του OCAI.",
            "Το OCAI οδηγεί σε ανάλογη εικόνα, καθώς αναδεικνύει κυρίως συνεργατική και καινοτομική κατεύθυνση, με σαφώς χαμηλότερη αγορακεντρική και ιεραρχική έμφαση. Αυτό δείχνει ότι η προσαρμοστικότητα δεν είναι περιφερειακό στοιχείο, αλλά ενσωματωμένο μέρος της πολιτισμικής λογικής της εταιρείας.",
        ),
        (
            "Σημαντικό στοιχείο προς αυτή την κατεύθυνση είναι και η πρόθεση",
            "Η πρόθεση για μεγαλύτερη αξιοποίηση της τεχνητής νοημοσύνης και ανάπτυξη εσωτερικών λύσεων κινείται προς αυτή την κατεύθυνση. Ωστόσο, τα ευρήματα δείχνουν περισσότερο οργανωσιακή ετοιμότητα για καινοτομία παρά ήδη συστηματοποιημένο μηχανισμό καινοτομίας.",
        ),
        (
            "Η κατεύθυνση αυτή συνδέεται και με τη σύγχρονη συζήτηση γύρω από την καινοτομία",
            "Στο ίδιο πλαίσιο, ο Sułkowski επισημαίνει ότι η καινοτομία διατηρείται όταν η ηγεσία ενισχύει πειραματισμό, διαλειτουργική συνεργασία και συνεχή βελτίωση, ενώ η τεχνητή νοημοσύνη μπορεί να λειτουργήσει όχι μόνο ως εργαλείο αποδοτικότητας αλλά και ως υποδομή καλύτερης μάθησης και λήψης αποφάσεων (Sułkowski and Lenart-Gansiniec, 2025, σελ. 159-161).",
        ),
        (
            "Η προσαρμοστικότητα, όμως, δεν αφορά μόνο την τεχνολογία.",
            "Η προσαρμοστικότητα, όμως, δεν αφορά μόνο την τεχνολογία. Αφορά και την ικανότητα του οργανισμού να εντάσσει γρήγορα νέους εργαζομένους, να ανακατανέμει γνώση και να ανταποκρίνεται χωρίς να χάνει ποιότητα εκτέλεσης.",
        ),
        (
            "Συνολικά, η Intragen φαίνεται να λειτουργεί σε ένα περιβάλλον",
            "Συνολικά, η Intragen λειτουργεί σε περιβάλλον όπου η καινοτομία είναι εφικτή, αλλά χρειάζεται πιο σταθερούς μηχανισμούς που να μετατρέπουν την τεχνική γνώση σε επαναλαμβανόμενη οργανωσιακή βελτίωση.",
        ),
        (
            "Από τα διαθέσιμα στοιχεία προκύπτει ότι η Intragen φαίνεται",
            "Τα διαθέσιμα στοιχεία δείχνουν ότι η Intragen επενδύει όχι μόνο στην εκπαίδευση, αλλά και στη δημιουργία περιβάλλοντος σύνδεσης με τον οργανισμό. Η έμφαση στους ανθρώπους, μαζί με τον προσανατολισμό στην εξυπηρέτηση πελάτη, υποδηλώνει προσπάθεια ισορροπίας ανάμεσα στην απόδοση και στην ανθρώπινη πλευρά της εργασίας.",
        ),
        (
            "Σε αυτό το πλαίσιο, φαίνεται ότι διαθέτει αρκετά στοιχεία",
            "Στα μη χρηματικά κίνητρα συμβάλλουν η συχνή κοινωνική αλληλεπίδραση, οι συναντήσεις ανά χώρα και η ετήσια συνάντηση, γιατί ενισχύουν γνωριμία, συνοχή και αίσθηση συμμετοχής.",
        ),
        (
            "Παράλληλα, η αναγνώριση της καλής απόδοσης αποτελεί ακόμη ένα θετικό στοιχείο.",
            "Η αναγνώριση της καλής απόδοσης αποτελεί ακόμη ένα θετικό στοιχείο. Οι διακρίσεις σε μηνιαία και τριμηνιαία βάση λειτουργούν ενισχυτικά όταν συνδυάζονται με δυνατότητες εξέλιξης και μάθησης. Αυτό ταιριάζει και με την προσέγγιση των Buchanan και Huczynski, οι οποίοι εντάσσουν στο πακέτο ανταμοιβών όχι μόνο τις οικονομικές απολαβές, αλλά και την αναγνώριση, τις ευκαιρίες ανάπτυξης και το εργασιακό περιβάλλον (Buchanan and Huczynski, 2013, σελ. 300).",
        ),
        (
            "Εξίσου σημαντικός φαίνεται να είναι και ο τρόπος ανατροφοδότησης",
            "Εξίσου σημαντικός είναι ο τρόπος ανατροφοδότησης και αξιολόγησης. Η εξαμηνιαία αξιολόγηση και οι μηνιαίες ατομικές συναντήσεις δείχνουν προσπάθεια πιο ολοκληρωμένης αποτίμησης και συνεχούς καθοδήγησης. Στο ίδιο πνεύμα, οι Buchanan and Huczynski τονίζουν ότι η ανατροφοδότηση από προϊσταμένους και συναδέλφους επηρεάζει καθοριστικά το πώς ο εργαζόμενος αντιλαμβάνεται την πορεία του μέσα στον οργανισμό (Buchanan and Huczynski, 2013, σελ. 316-317).",
        ),
        (
            "Σημαντικό ρόλο στα κίνητρα φαίνεται να παίζει και το οικονομικό σκέλος.",
            "Ρόλο παίζει και το οικονομικό σκέλος, αφού οι αμοιβές θεωρούνται καλές για τα ελληνικά δεδομένα. Ωστόσο, τα ευρήματα δείχνουν ότι η δέσμευση δεν στηρίζεται μόνο στον μισθό, αλλά και στην αναγνώριση, στις σχέσεις, στην ανατροφοδότηση και στην αίσθηση εξέλιξης.",
        ),
        (
            "Η παραπάνω εικόνα συνδέεται και με τα ευρήματα του OCAI, τα οποία δείχνουν",
            "Το OCAI ενισχύει αυτή την ερμηνεία, καθώς δείχνει συνεργατική και ανθρωποκεντρική κουλτούρα. Στο παρόν κείμενο το εύρημα αξιοποιείται κυρίως ως ένδειξη των κυρίαρχων πολιτισμικών τάσεων και όχι ως εξαντλητική χαρτογράφηση όλων των εμπειριών.",
        ),
        (
            "Παρότι η συνολική εικόνα είναι θετική, προκύπτει και εδώ ένα σημείο προσοχής.",
            "Παρότι η συνολική εικόνα είναι θετική, προκύπτει και εδώ ένα κρίσιμο σημείο. Η θετική κουλτούρα δεν αρκεί από μόνη της για διατηρήσιμη δέσμευση, αν δεν συνοδεύεται από σαφή αίσθηση προόδου και ορατά επόμενα βήματα εξέλιξης.",
        ),
        (
            "Συνολικά, η Intragen φαίνεται να διαθέτει αρκετούς παράγοντες",
            "Συνολικά, η Intragen διαθέτει αρκετούς παράγοντες που ενισχύουν κίνητρα και δέσμευση. Η κρίσιμη πρόκληση είναι να διατηρηθούν όσο ο οργανισμός μεγαλώνει και να συνδεθεί ακόμη πιο καθαρά η θετική εργασιακή εμπειρία με ουσιαστική επαγγελματική εξέλιξη.",
        ),
        (
            "Μια πρώτη κατεύθυνση βελτίωσης αφορά τη δημιουργία μιας μικρής ομάδας",
            "Μια πρώτη κατεύθυνση βελτίωσης αφορά τη δημιουργία μικρής ομάδας ανάπτυξης εσωτερικών εφαρμογών και εργαλείων παραγωγικότητας. Στόχος δεν είναι μια γενική πρωτοβουλία καινοτομίας, αλλά ένας συγκεκριμένος μηχανισμός που θα μετατρέπει επαναλαμβανόμενα προβλήματα της καθημερινής λειτουργίας σε εφαρμόσιμες λύσεις.",
        ),
        (
            "Η ομάδα αυτή θα μπορούσε να λειτουργεί σε άμεση επικοινωνία",
            "Η ομάδα αυτή θα μπορούσε να λειτουργεί σε άμεση επικοινωνία με τον τεχνικό διευθυντή, εντοπίζοντας επαναλαμβανόμενες διαδικασίες και σημεία καθυστέρησης. Η αξία της θα είναι μεγαλύτερη αν οι λύσεις βασίζονται σε πραγματικές ανάγκες διαφορετικών τμημάτων και όχι σε θεωρητικές υποθέσεις.",
        ),
        (
            "Η πρωτοβουλία αυτή συνδέεται άμεσα με την επιθυμία για μεγαλύτερη αξιοποίηση",
            "Η πρωτοβουλία αυτή συνδέεται με την επιθυμία για μεγαλύτερη αξιοποίηση της τεχνητής νοημοσύνης και με πιο πρακτική κατεύθυνση καινοτομίας. Η ανάπτυξη εσωτερικών εργαλείων μπορεί να μειώσει τον χρόνο που δαπανάται σε επαναλαμβανόμενες διαδικασίες και να απελευθερώσει χρόνο για εργασία υψηλότερης αξίας.",
        ),
        (
            "Επιπλέον, αυτή η κίνηση μπορεί να ενισχύσει και την ομαδικότητα.",
            "Η παρέμβαση μπορεί να ενισχύσει και την ομαδικότητα, αν οι εργαζόμενοι συμμετέχουν στον εντοπισμό αναγκών και στην παραγωγή ιδεών. Τότε η καινοτομία παύει να είναι αποσπασματική και γίνεται συλλογική οργανωσιακή διαδικασία.",
        ),
        (
            "Συνολικά, η δημιουργία μιας τέτοιας ομάδας μπορεί να βοηθήσει",
            "Συνολικά, μια τέτοια ομάδα μπορεί να μετατρέψει τη διάσπαρτη τεχνογνωσία σε πιο σταθερό μηχανισμό οργανωσιακής βελτίωσης.",
        ),
        (
            "Μια δεύτερη κατεύθυνση βελτίωσης αφορά τη δημιουργία μιας πιο δομημένης διαδικασίας",
            "Μια δεύτερη κατεύθυνση βελτίωσης αφορά πιο δομημένη μετάβαση από την εκπαίδευση στην ενεργή συμμετοχή σε έργα. Από την ανάλυση προκύπτει ότι η Intragen επενδύει ουσιαστικά στην αρχική εκπαίδευση, αλλά χρειάζεται πιο ομαλή μετάφρασή της σε πραγματική εμπειρία έργου.",
        ),
        (
            "Η λογική εδώ δεν είναι να αλλάξει το εκπαιδευτικό πλαίσιο",
            "Η λογική δεν είναι να αλλάξει το εκπαιδευτικό πλαίσιο, αλλά να οργανωθεί καλύτερα το επόμενο στάδιο. Η νωρίτερη παρακολούθηση έμπειρων συναδέλφων και η σταδιακή έκθεση σε πραγματικά έργα θα βοηθούσαν τους νέους εργαζομένους να συνδέουν γρηγορότερα την εκπαίδευση με τις πραγματικές απαιτήσεις του ρόλου.",
        ),
        (
            "Παράλληλα, αυτή η μετάβαση θα μπορούσε να οργανωθεί",
            "Η μετάβαση αυτή θα μπορούσε να οργανωθεί σε πιο καθαρά στάδια: παρατήρηση, ανάληψη μικρότερων καθηκόντων, συμμετοχή σε σαφή παραδοτέα και σταδιακή αύξηση ευθύνης. Έτσι, η ένταξη στην πράξη θα ακολουθεί πιο ορατή και σταθερή πορεία.",
        ),
        (
            "Στην ίδια κατεύθυνση, θα ήταν χρήσιμο να επανενεργοποιηθούν",
            "Θα ήταν επίσης χρήσιμο να επανενεργοποιηθούν πιο συστηματικά τα εργαστήρια ανά δύο εβδομάδες και τα βίντεο ανταλλαγής γνώσης, τα οποία φαίνεται ότι περιορίστηκαν λόγω αυξημένου φόρτου και ταχείας ανάπτυξης. Η επαναφορά τους θα ενίσχυε τη διάχυση γνώσης και τη σύνδεση της εμπειρίας των πιο έμπειρων στελεχών με την ανάπτυξη των νεότερων.",
        ),
        (
            "Η παρέμβαση αυτή είναι σημαντική, γιατί η πρακτική εμπειρία",
            "Η παρέμβαση αυτή είναι σημαντική, γιατί η πρακτική εμπειρία δεν ενισχύει μόνο την τεχνική ανάπτυξη αλλά και την αυτοπεποίθηση, την αίσθηση προόδου και την επαγγελματική ωρίμανση των νεότερων εργαζομένων.",
        ),
        (
            "Η κατεύθυνση αυτή αξιοποιεί ουσιαστικά τη βάση που ήδη υπάρχει",
            "Έτσι, η ισχυρή εκπαιδευτική βάση της εταιρείας θα συνδέεται πιο καθαρά με προγραμματισμένη ανάληψη ρόλου και πιο προβλέψιμη στελέχωση έργων.",
        ),
        (
            "Μια τρίτη κατεύθυνση βελτίωσης αφορά την πιο πρακτική αξιοποίηση",
            "Μια τρίτη κατεύθυνση βελτίωσης αφορά την πιο πρακτική αξιοποίηση του ρόλου του επικεφαλής ομάδας στην καθημερινή ανάπτυξη των εργαζομένων. Η βελτίωση δεν απαιτεί νέο σύστημα, αλλά πιο συστηματική οργάνωση πρακτικών που ήδη υπάρχουν.",
        ),
        (
            "Πιο συγκεκριμένα, θα μπορούσε να καθιερωθεί για κάθε εργαζόμενο",
            "Θα μπορούσε να καθιερωθεί για κάθε εργαζόμενο ένα σύντομο ατομικό πλάνο εξέλιξης, το οποίο θα ενημερώνεται στις μηνιαίες συναντήσεις με τον επικεφαλής ομάδας. Το πλάνο δεν χρειάζεται να είναι γραφειοκρατικό: αρκεί να περιλαμβάνει τον βασικό στόχο του επόμενου διαστήματος, λίγα συγκεκριμένα βήματα και μια σύντομη αποτύπωση δυνατών σημείων και αναγκών ανάπτυξης.",
        ),
        (
            "Παράλληλα, η εξαμηνιαία αξιολόγηση θα μπορούσε να συνδεθεί",
            "Η εξαμηνιαία αξιολόγηση θα μπορούσε να συνδεθεί πιο άμεσα με αυτό το πλάνο. Αντί να λειτουργεί ως ξεχωριστή αποτίμηση, θα μπορούσε να εξετάζει πρόοδο, εμπόδια και επόμενα βήματα, ώστε να αποκτά σαφέστερο αναπτυξιακό χαρακτήρα.",
        ),
        (
            "Μια ακόμη πρακτική βελτίωση θα ήταν να υπάρχει πιο σταθερή δομή",
            "Μια ακόμη βελτίωση θα ήταν πιο σταθερή δομή στην ανατροφοδότηση από άλλους εργαζομένους. Αντί για γενικές εντυπώσεις, η συμβολή τους θα μπορούσε να βασίζεται σε λίγα σταθερά ερωτήματα, ώστε η αξιολόγηση να γίνεται καθαρότερη και πιο χρήσιμη χωρίς να γίνεται πιο περίπλοκη.",
        ),
        (
            "Η κατεύθυνση αυτή είναι ιδιαίτερα σημαντική για νεότερους εργαζομένους",
            "Η κατεύθυνση αυτή είναι ιδιαίτερα σημαντική για νεότερους εργαζομένους, οι οποίοι χρειάζονται πιο σαφή εικόνα για το πού βρίσκονται και πώς μπορούν να εξελιχθούν μέσα στην εταιρεία. Ταυτόχρονα, ενισχύει τη δέσμευση, γιατί κάνει πιο ορατή την επένδυση της εταιρείας στην πρόοδό τους.",
        ),
        (
            "Συνολικά, η βελτίωση αυτή δεν απαιτεί ριζική αλλαγή",
            "Συνολικά, η βελτίωση αυτή αξιοποιεί υπάρχουσες πρακτικές και τις μετατρέπει σε πιο σταθερό μηχανισμό ανάπτυξης και καθοδήγησης.",
        ),
        (
            "Η ανάλυση της Intragen έδειξε ότι πρόκειται για έναν οργανισμό",
            "Η ανάλυση της Intragen έδειξε οργανισμό που βασίζεται σε υψηλή εξειδίκευση, έργο-κεντρική λειτουργία και συνεργατική κουλτούρα. Η συνολική εικόνα συνδυάζει τεχνική ποιότητα με ανάγκη για συνεχή συντονισμό και προσαρμογή.",
        ),
        (
            "Τα ευρήματα των ερωτηματολογίων ενίσχυσαν αυτή την εικόνα.",
            "Τα ευρήματα των ερωτηματολογίων ενίσχυσαν αυτή την εικόνα. Ο Mintzberg έδειξε κυρίως επαγγελματική γραφειοκρατία με στοιχεία ευελιξίας, ενώ το OCAI ανέδειξε συνεργατική και ανθρωποκεντρική κουλτούρα με ορισμένα στοιχεία καινοτομικής κατεύθυνσης. Επειδή, όμως, η πρωτογενής έρευνα βασίστηκε σε μία ενημερωμένη εσωτερική οπτική, τα συμπεράσματα διαβάζονται πιο σωστά ως ισχυρές διαγνωστικές τάσεις παρά ως πλήρης οργανωσιακή γενίκευση.",
        ),
        (
            "Η ανάλυση ανέδειξε αρκετά δυνατά σημεία, όπως η εξάμηνη εκπαίδευση",
            "Η ανάλυση ανέδειξε δυνατά σημεία, όπως η εξάμηνη εκπαίδευση, η καθοδήγηση από μέντορα και επικεφαλής ομάδας, η έμφαση στη συνεργασία και η προσπάθεια διατήρησης υποστηρικτικού εργασιακού περιβάλλοντος. Ταυτόχρονα, έδειξε ότι το βασικό πεδίο βελτίωσης είναι η ταχύτερη μετατροπή της μάθησης σε συμμετοχή σε έργα, η πιο σταθερή αξιοποίηση της καινοτομίας και η περαιτέρω ενίσχυση του αναπτυξιακού ρόλου της ηγεσίας.",
        ),
        (
            "Με βάση τα παραπάνω, οι προτάσεις βελτίωσης εστίασαν",
            "Με βάση τα παραπάνω, οι προτάσεις βελτίωσης εστίασαν στην ανάπτυξη εσωτερικών εργαλείων παραγωγικότητας, σε πιο δομημένη μετάβαση από την εκπαίδευση στην πράξη και στην καλύτερη αξιοποίηση του ρόλου του επικεφαλής ομάδας και της αξιολόγησης. Η Intragen φαίνεται να διαθέτει ισχυρές βάσεις για περαιτέρω ανάπτυξη, εφόσον συνδέσει πιο αποτελεσματικά τη μάθηση, την καινοτομία και την εξέλιξη των ανθρώπων της.",
        ),
    ]
    for prefix, new_text in paragraph_updates:
        set_paragraph_text_by_prefix(doc, prefix, new_text, style_name="Normal")

    insert_after_prefix(
        doc,
        "Τα δύο συμπληρωμένα ερωτηματολόγια αντιμετωπίστηκαν ως πρωτογενές υλικό",
        [
            "Η επιλογή αυτή προσφέρει επαρκές βάθος για διάγνωση, επειδή ο συμμετέχων έχει άμεση εικόνα διαδικασιών ένταξης, καθημερινής συνεργασίας και ανατροφοδότησης. Παράλληλα, λειτουργίες που σχετίζονται περισσότερο με εμπορικές, διοικητικές ή υποστηρικτικές δραστηριότητες ενδέχεται να αποτυπώνονται ασθενέστερα, γι' αυτό τα ευρήματα ερμηνεύονται με προσοχή.",
        ],
    )
    insert_after_prefix(
        doc,
        "Η έμφαση στην άτυπη υποστήριξη, στην εύκολη πρόσβαση σε καθοδήγηση",
        [
            "Παράλληλα, η εταιρεία φαίνεται να συνδυάζει ευελιξία με σαφείς μηχανισμούς συντονισμού, όπως καταγραφή χρόνου, δηλωμένη τοποθεσία εργασίας και επαναλαμβανόμενες ομαδικές συναντήσεις. Αυτό δείχνει ότι η καθημερινή λειτουργία δεν αφήνεται αποκλειστικά στην άτυπη συνεργασία, αλλά πλαισιώνεται από συγκεκριμένες ρουτίνες παρακολούθησης και ευθυγράμμισης.",
        ],
    )
    insert_after_prefix(
        doc,
        "Το βασικό εύρημα του Mintzberg είναι ότι η Intragen δεν εντάσσεται καθαρά",
        [
            "Η υβριδική αυτή εικόνα έχει ιδιαίτερη σημασία, γιατί εξηγεί γιατί ο οργανισμός δείχνει ταυτόχρονα ανάγκη για σταθερότητα και ανάγκη για προσαρμογή. Αν η ισορροπία γείρει υπερβολικά προς την εξειδίκευση, μπορεί να δυσκολευτεί η ταχύτητα απόκρισης. Αν, αντίθετα, γείρει υπερβολικά προς την ευελιξία, μπορεί να αποδυναμωθεί η τυπικότητα γνώσης που απαιτείται σε έργα IAM.",
        ],
    )
    insert_after_prefix(
        doc,
        "Το εύρημα του Mintzberg ενισχύει αυτή την εικόνα, αλλά αναδεικνύει",
        [
            "Από διοικητική σκοπιά, αυτό σημαίνει ότι η Intragen φαίνεται να έχει ήδη επενδύσει σωστά στην ποιότητα της τεχνικής προετοιμασίας, αλλά δεν έχει ακόμη κλείσει πλήρως τον κύκλο από τη μάθηση στην απόδοση. Όσο η γνώση παραμένει έντονα συνδεδεμένη με μέντορες και έμπειρα στελέχη, τόσο μεγαλύτερος γίνεται ο κίνδυνος η ένταξη να επιβραδύνεται όταν αυξάνεται ο αριθμός των νεοπροσλαμβανόμενων. Άρα, η πρόκληση δεν είναι περισσότερη εκπαίδευση, αλλά καλύτερη μεταφορά της ήδη υπάρχουσας γνώσης σε παραγωγική συμμετοχή.",
        ],
    )
    insert_after_prefix(
        doc,
        "Παράλληλα, η μάθηση δεν περιορίζεται στην τυπική εκπαίδευση.",
        [
            "Σε ένα περιβάλλον IAM, η μάθηση φαίνεται να περιλαμβάνει και διάσταση ασφάλειας πληροφοριών και υπεύθυνης διαχείρισης δεδομένων. Αυτό σημαίνει ότι η διαδικασία ένταξης δεν αφορά μόνο τεχνικά εργαλεία και έργα, αλλά και εξοικείωση με απαιτήσεις απορρήτου, ασφαλούς ανάπτυξης λογισμικού, απομακρυσμένης πρόσβασης και πειθαρχημένης χρήσης πληροφοριακών πόρων.",
        ],
    )
    insert_after_prefix(
        doc,
        "Συνολικά, η Intragen λειτουργεί σε περιβάλλον όπου η καινοτομία είναι εφικτή",
        [
            "Και εδώ η κριτική ανάγνωση είναι σημαντική. Το γεγονός ότι ο οργανισμός εμφανίζει στοιχεία καινοτομίας δεν σημαίνει ότι διαθέτει ήδη ώριμο σύστημα παραγωγής και κλιμάκωσης νέων ιδεών. Τα πρωτογενή δεδομένα δείχνουν περισσότερο βούληση προσαρμογής παρά πλήρως θεσμοποιημένη διακυβέρνηση καινοτομίας: υπάρχουν τεχνολογική ευαισθησία, συνεργασία και πρόθεση, αλλά λιγότερο σαφείς ρουτίνες για προτεραιοποίηση, πειραματισμό και διάχυση λύσεων σε όλο τον οργανισμό.",
        ],
    )
    insert_after_prefix(
        doc,
        "Συνολικά, η Intragen διαθέτει αρκετούς παράγοντες που ενισχύουν κίνητρα",
        [
            "Ο βασικός κίνδυνος εντοπίζεται λιγότερο στο σημερινό κλίμα και περισσότερο στην κλιμάκωση: αν οι μηχανισμοί ανατροφοδότησης, αξιολόγησης και εξέλιξης δεν γίνουν πιο σταθεροί, η σημερινή θετική εμπειρία μπορεί να παραμείνει αποσπασματική και όχι διατηρήσιμη.",
        ],
    )
    insert_after_prefix(
        doc,
        "Στα μη χρηματικά κίνητρα συμβάλλουν",
        [
            "Η εικόνα αυτή ενισχύεται όταν η οργανωσιακή φροντίδα συνδέεται και με πιο δομημένες παροχές, όπως υποστήριξη ευεξίας, ανάπτυξη μέσω εκπαιδευτικών ευκαιριών και συμβολικές μορφές αναγνώρισης. Τέτοιες πρακτικές κάνουν την εμπειρία του εργαζομένου πιο συνεπή και περιορίζουν το ενδεχόμενο η δέσμευση να εξαρτάται μόνο από το άμεσο κλίμα της ομάδας.",
        ],
    )
    insert_after_prefix(
        doc,
        "Συνολικά, μια τέτοια ομάδα μπορεί να μετατρέψει τη διάσπαρτη τεχνογνωσία",
        [
            "Η διοικητική της αξία βρίσκεται κυρίως στη διακυβέρνηση της προσπάθειας: σαφή κριτήρια επιλογής θεμάτων, μικρό αλλά προστατευμένο χρόνο ενασχόλησης και βασικούς δείκτες αποτελέσματος, ώστε η καινοτομία να μη χαθεί μέσα στις καθημερινές πιέσεις της εκτέλεσης έργων. Ένα σύντομο πιλοτικό κύμα εφαρμογών με 2-3 στοχευμένες περιπτώσεις χρήσης θα επέτρεπε στην εταιρεία να μετρήσει γρήγορα υιοθέτηση, εξοικονόμηση χρόνου και πραγματική οργανωσιακή αξία.",
        ],
    )
    insert_after_prefix(
        doc,
        "Έτσι, η ισχυρή εκπαιδευτική βάση της εταιρείας θα συνδέεται πιο καθαρά",
        [
            "Η πρόταση αυτή απαντά στο πιο εμφανές κενό της ανάλυσης: την έλλειψη καθαρού ενδιάμεσου σταδίου ανάμεσα στην εκπαίδευση και στην πλήρη παραγωγική εμπλοκή. Αν το σημείο αυτό μείνει άτυπο, αυξάνονται η αβεβαιότητα, η εξάρτηση από έμπειρους συναδέλφους και το κόστος προσαρμογής. Αν οργανωθεί καλύτερα, βελτιώνεται ταυτόχρονα η εμπειρία των νεότερων στελεχών και η επιχειρησιακή συνέπεια.",
        ],
    )
    insert_after_prefix(
        doc,
        "Συνολικά, η βελτίωση αυτή αξιοποιεί υπάρχουσες πρακτικές",
        [
            "Αυτή η κατεύθυνση δίνει αξία γιατί καθιστά την αξιολόγηση πιο προβλέψιμο εργαλείο λήψης αποφάσεων για ανάπτυξη, κατανομή ευθυνών και διατήρηση ικανών εργαζομένων. Έτσι, ο επικεφαλής ομάδας λειτουργεί λιγότερο ως αποσπασματικός αξιολογητής και περισσότερο ως σταθερός σύνδεσμος ανάμεσα στις οργανωσιακές προσδοκίες και στην ατομική πρόοδο. Παράλληλα, δίνεται καλύτερη εικόνα και στη διοίκηση για το πότε ένας εργαζόμενος είναι έτοιμος για μεγαλύτερη ευθύνη ή για πιο απαιτητικά έργα.",
        ],
    )
    insert_after_prefix(
        doc,
        "Η πρόσφατη βιβλιογραφία ενισχύει αυτή την κατεύθυνση, δείχνοντας",
        [
            "Από διοικητική σκοπιά, το σημαντικότερο συμπέρασμα δεν είναι ότι η Intragen χρειάζεται ριζική αναδιοργάνωση, αλλά ότι χρειάζεται να ωριμάσει οργανωσιακά γύρω από ήδη θετικά στοιχεία. Η εταιρεία διαθέτει ποιότητα γνώσης, συνεργατική κουλτούρα και πρόθεση καινοτομίας. Εκεί όπου χρειάζεται ενίσχυση είναι η συνέπεια των μηχανισμών που μετατρέπουν αυτά τα πλεονεκτήματα σε επαναλαμβανόμενο αποτέλεσμα: ταχύτερη ένταξη σε έργα, πιο ορατή εξέλιξη ρόλου, πιο δομημένη ανατροφοδότηση και καλύτερη μεταφορά γνώσης από τους έμπειρους στους νεότερους. Με αυτή τη λογική, οι παρεμβάσεις δεν χρειάζεται να εφαρμοστούν όλες ταυτόχρονα: μπορούν να ακολουθήσουν αλληλουχία ωρίμανσης, ξεκινώντας από τη μετάβαση στην πράξη, περνώντας στη συστηματικότερη καθοδήγηση και καταλήγοντας σε πιο οργανωμένη αξιοποίηση της εσωτερικής καινοτομίας.",
        ],
    )

    replace_page_refs_in_body(doc)
    restore_body_figure_1(doc)
    replace_ocai_profile_figures(doc)
    replace_ocai_dimension_figures(doc)
    enhance_appendices(doc)
    insert_ocai_comparison_table(doc)
    polish_cover_page(doc)
    sync_manual_contents(doc)
    normalize_text_paragraphs(doc)
    apply_body_spacing(doc)

    saved_path = save_document(doc)
    print(f"Saved revised document to: {saved_path}")


if __name__ == "__main__":
    main()
